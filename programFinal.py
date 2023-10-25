import json
import base64
import win32com.client
import win32com.client as win32
import re
import win32api
import urllib.parse  # Importado para analizar la URL
import requests
import os
import sys
import subprocess
from docxtpl import DocxTemplate
from docx import Document
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
import docx
import time
from collections import OrderedDict

# Crear un entorno virtual llamado "shell"
subprocess.run([sys.executable, "-m", "venv", "shell"])

# Definir la ubicación del ejecutable de Python en el entorno virtual
venv_python = os.path.join("shell", "Scripts", "python")
if sys.platform == "linux":
    venv_python = os.path.join("shell", "bin", "python")


# Actualizar pip en el entorno virtual
subprocess.run([venv_python, "-m", "pip", "install", "--upgrade", "pip"])


# Instalar las dependencias especificadas en el archivo requirements.txt
subprocess.run([venv_python, "-m", "pip", "install", "-r", "requirements.txt"])

# Pregunta al usuario si quiere descargar todas las páginas y subpáginas de la Wiki
respuesta = input("¿Quieres descargar todas las páginas y subpáginas de la Wiki? (s/n): ")

if respuesta.lower() == 's':
    def todaslaspaginas():
        import requests
        import json
        import base64
        import re
        from docxtpl import DocxTemplate
        import os
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsmap
        from collections import OrderedDict
        import docx
        import win32com.client
        import time


        def sanitize_placeholder(placeholder):
            # Replace spaces and invalid characters with underscores
            sanitized = re.sub(r'[^\w]', '_', placeholder)
            # Remove leading digits and underscores to ensure a valid variable name
            sanitized = re.sub(r'^\d+|_', '', sanitized)
            return sanitized


        def update_toc(docx_file):
            try:
                word = win32com.client.DispatchEx("Word.Application")
                doc = word.Documents.Open(docx_file)
                doc.TablesOfContents(1).Update()
                doc.Close(SaveChanges=True)
                word.Quit()
            except Exception as e:
                print(f"An error occurred while updating the table of contents: {e}")



        def get_page_content(url):
            content_url = url + "?api-version=7.0&includeContent=true"
            response = requests.get(content_url, headers=headers)

            if response.status_code == 200:
                return json.loads(response.text)['content']
            else:
                print(f"Error al obtener el contenido de la página: {response.status_code}")
                return "Contenido no disponible"

        def extract_pages_recursive(page, level=1):
            if not page:
                return []

            original_content = get_page_content(page['url'])
            content = get_page_content(page['url'])
            # Replace <span> tags with plain text
            content = re.sub(r'<span style="color:[^>]*>([^<]*)</span>', r'\1', content)
            # Reemplazar <b><span style="color:blue">Note:</span></b> con 'Note:' de manera dinámica
            content = re.sub(r'<b>([^<]*)</b>', r'\1', content, flags=re.IGNORECASE)
            # Reemplazar <span style="color:(.*?)">(.*?)</span> con '(.*?)'
            content = re.sub(r'<span style="color:(.*?)">(.*?)</span>', r'\2', content)
            # Reemplazar <center>(.*?)</center> con '(.*?)'
            content = re.sub(r'<center>(.*?)</center>', r'\1', content)
            # Reemplazar <code>(.*?)</code> con '(.*?)'
            content = re.sub(r'<code>(.*?)</code>', r'\1', content)
            # Reemplazar <br>(.*?)</br> con '(.*?)'
            content = re.sub(r'<br>(.*?)</br>', r'\1', content)
            # Captura el color y el texto hasta la siguiente etiqueta o el final de la línea.
            content = re.sub(r'<span style="color:(.*?)">([^<]*)', r'\2', content)
            # reemplaza etiqueta <br> con un salto de linea
            content = re.sub(r'<br>', '\n', content)
            # Reemplazar <Lista> @<lo que sea> con Lista @lo que sea
            content = re.sub(r'<Lista> @<([^>]+)>', r'Lista @\1', content)
            
            


            info = {
                'name': page['path'],
                'short_name': page['path'].split('/')[-1],
                'url': page['url'],
                'original_content': original_content,
                'content': content ,
                'level': level,
                'subpages': []
            }


            page_info = [info]

            if 'subPages' in page:
                for sub_page in page['subPages']:
                    page_info.extend(extract_pages_recursive(sub_page, level+1))

            return page_info

        def extract_url_values(url):
            regex = r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/.*"
            match = re.search(regex, url)
            if match:
                return match.groupdict()
            else:
                return None

        def extract_placeholders(template_path):
            doc = Document(template_path)
            placeholders = []

            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        
                        # Find placeholders in the full text
                        start_index = full_text.find('{{')
                        end_index = full_text.find('}}')
                        while start_index != -1 and end_index != -1:
                            placeholder = full_text[start_index+2:end_index].strip()
                            placeholders.append(placeholder)
                            
                            # Find the next placeholder
                            start_index = full_text.find('{{', end_index)
                            end_index = full_text.find('}}', end_index+2)

                elif element.tag.endswith('tc'):  # Check for table cell (td)
                    for p in element.iterchildren('{%s}p' % nsmap['w']):
                        paragraph = Paragraph(p, doc)
                        if hasattr(paragraph, 'runs'):
                            # Concatenate the text of adjacent runs
                            full_text = ''.join([run.text for run in paragraph.runs])
                            
                            # Find placeholders in the full text
                            start_index = full_text.find('{{')
                            end_index = full_text.find('}}')
                            while start_index != -1 and end_index != -1:
                                placeholder = full_text[start_index+2:end_index].strip()
                                placeholders.append(placeholder)
                                
                                # Find the next placeholder
                                start_index = full_text.find('{{', end_index)
                                end_index = full_text.find('}}', end_index+2)

            # Remove duplicates from the list of placeholders while maintaining the order of the elements
            placeholders = list(OrderedDict.fromkeys(placeholders))

            return placeholders


        def create_context(page_info, placeholders):
            context = {}
            title_index = 1
            for placeholder in placeholders:
                if title_index < len(page_info):
                    page = page_info[title_index]
                    # Agregar valores al contexto tanto para marcadores de posición de título como para marcadores de posición de contenido
                    if placeholder.endswith('_content'):
                        if page['content'] != 'No hay contenido':
                            context[placeholder] = page['content']
                        title_index += 1
                    else:
                        context[placeholder] = page['short_name']
            return context


        wiki_url = input("Introduce la URL principal del portal Wiki de Azure: ")
        url_values = extract_url_values(wiki_url)

        if url_values is None:
            print("La URL proporcionada no es válida.")
            exit()

        organization = url_values['organization']
        project = url_values['project']
        wiki = url_values['wiki']

        personal_access_token = input("Introduce tu token de acceso personal: ")

        credentials = f":{personal_access_token}"
        encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')

        headers = {
            'Authorization': f'Basic {encoded_credentials}',
            "Content-Type": "application/json",
            "Accept": "application/json"
        }

        url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki}/pages?api-version=7.0&recursionLevel=full"

        response = requests.get(url, headers=headers)

        if response.status_code == 200:
            root_page = json.loads(response.text)
        else:
            print(f"Error al obtener la página raíz de la wiki: {response.status_code}")
            root_page = {}

        page_info = extract_pages_recursive(root_page)


        # Save the .md file with original content
        md_filename_original = 'htmlymd.md'
        with open(md_filename_original, 'w', encoding='utf-8') as f:
            for page in page_info:
                # Get the title and original content of the page
                title = page['name'].split('/')[-1]
                original_content = page['original_content'].strip()

                # Only write the title and original content to the Markdown file if the original content is not empty
                if original_content and not (title.startswith("#") and "No hay contenido" in title) and page['name'] != "/":
                    f.write(f'# {title}\n')
                    f.write(original_content)
                    f.write('\n\n')


        # Save the .md file
        md_filename = 'todosmd.md'
        with open('todosmd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Get the title and content of the page
                title = page['name'].split('/')[-1]
                content = page['content'].strip()

                # Only write the title and content to the Markdown file if the content is not empty
                if content and not (title.startswith("#") and "No hay contenido" in title) and page['name'] != "/":
                    f.write(f'# {title}\n')
                    f.write(content)
                    f.write('\n\n')


        # Read the Markdown content
        with open(md_filename, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        # Extract the level 1 headings with regex
        title_pattern = r'^# (.*)'
        titles = re.findall(title_pattern, markdown_content, flags=re.M)


        # Generate the placeholders
        placeholders = [{'text': f'Titulo{i+1}', 'level': page['level']} for i, page in enumerate(page_info)]

        def add_titles_to_template(template_path, titles_info):
            doc = Document(template_path)

            # Find the element after which to insert the placeholders
            insert_element = None
            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if 'REFERENCES' in full_text:
                            insert_element = paragraph
                            break

            if insert_element is not None:
                # Insert the titles before the element with the desired heading style
                for title_info in titles_info:
                    # Skip titles with name '/' or empty short_name
                    if 'name' in title_info and (title_info['name'] == '/' or not title_info['short_name']):
                        continue
                    sanitized_title = sanitize_placeholder(title_info["title"])
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}}}}}')
                    # Calculate the maximum level in the JSON
                    max_level = max([title_info['level'] for title_info in titles_info])
                    # Generate the level_map dictionary dynamically
                    level_map = {i: f'Heading {i}' for i in range(1, max_level+1)}
                    heading_style = level_map.get(title_info["level"] - 1, 'Normal')
                    p.style = heading_style
                    # Insert a placeholder for the content below the title
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}_content}}}}')

                # Insert a page break after the last title
                p = insert_element.insert_paragraph_before()
                run = p.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)

            # Save the modified template
            doc.save('new_template.docx')


        # Generate the title information
        titles_info = [{'title': page['name'].split('/')[-1], 'level': page['level']} for page in page_info if page['content'].strip() and page['name'].split('/')[-1].strip()]


        # Call the function to create a new template and add the titles to it
        add_titles_to_template('plantilla.docx', titles_info)

        # Load the modified template
        doc = DocxTemplate("new_template.docx")

        # Extract the placeholders from the template
        placeholders = extract_placeholders("new_template.docx")


        # Create a new context
        context = create_context(page_info, placeholders)


        # Render template with dynamic context
        doc.render(context)

        # Save the generated document
        doc.save("documento_generado.docx")


        print("Wait one moment to finish work")
        # Wait for a few seconds to make sure the file is saved
        time.sleep(4)
        print("Ya, todo ready by GPT")

        # Check if the file was saved successfully
        if os.path.exists("documento_generado.docx"):
            print("El archivo documento_generado.docx se ha guardado correctamente.")
        else:
            print("No se pudo guardar el archivo documento_generado.docx.")

        # Get the directory of the script
        dir_path = os.path.dirname(os.path.realpath(__file__))

        # Build the file path
        docx_file = os.path.join(dir_path, "documento_generado.docx")

        # Call the function to update the table of contents in the generated document
        update_toc(docx_file)

        
        pass
        
    todaslaspaginas()  # Llama a la función todaslaspaginas

else:

    def paginaconcreta2():
        import requests
        import json
        import base64
        import re
        from docxtpl import DocxTemplate
        import os
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsmap
        from collections import OrderedDict
        import docx
        import win32com.client
        import time


        def sanitize_placeholder(placeholder):
            # Replace spaces and invalid characters with underscores
            sanitized = re.sub(r'[^\w]', '_', placeholder)
            # Remove leading digits and underscores to ensure a valid variable name
            sanitized = re.sub(r'^\d+|_', '', sanitized)
            return sanitized

        def update_toc(docx_file):
            word = win32com.client.DispatchEx("Word.Application")
            doc = word.Documents.Open(docx_file)
            doc.TablesOfContents(1).Update()
            doc.Close(SaveChanges=True)
            word.Quit()

        def get_page_content(url, headers):
            content_url = url + "?api-version=7.0&includeContent=true"
            response = requests.get(content_url, headers=headers)

            if response.status_code == 200:
                # Convierte la respuesta en un objeto JSON
                response_json = json.loads(response.text)
                # Devuelve el valor del campo 'content'
                return response_json['content']
            else:
                print(f"Error al obtener el contenido de la página: {response.status_code}")
                return "Contenido no disponible"

        def extract_pages_recursive(page, headers=None, level=1):
            if not page:
                return []

            original_content = get_page_content(page['url'], headers)
            content = get_page_content(page['url'], headers)
            # Replace <span> tags with plain text
            content = re.sub(r'<span style="color:[^>]*>([^<]*)</span>', r'\1', content)
            # Reemplazar <b><span style="color:blue">Note:</span></b> con 'Note:' de manera dinámica
            content = re.sub(r'<b>([^<]*)</b>', r'\1', content, flags=re.IGNORECASE)
            # Reemplazar <span style="color:(.*?)">(.*?)</span> con '(.*?)'
            content = re.sub(r'<span style="color:(.*?)">(.*?)</span>', r'\2', content)
            # Reemplazar <center>(.*?)</center> con '(.*?)'
            content = re.sub(r'<center>(.*?)</center>', r'\1', content)
            # Reemplazar <code>(.*?)</code> con '(.*?)'
            content = re.sub(r'<code>(.*?)</code>', r'\1', content)
            # Reemplazar <br>(.*?)</br> con '(.*?)'
            content = re.sub(r'<br>(.*?)</br>', r'\1', content)
            # Captura el color y el texto hasta la siguiente etiqueta o el final de la línea.
            content = re.sub(r'<span style="color:(.*?)">([^<]*)', r'\2', content)
            # reemplaza etiqueta <br> con un salto de linea
            content = re.sub(r'<br>', '\n', content)
            # Reemplazar <Lista> @<lo que sea> con Lista @lo que sea
            content = re.sub(r'<Lista> @<([^>]+)>', r'Lista @\1', content)
            


            info = {
                'name': page['path'],
                'short_name': page['path'].split('/')[-1],
                'url': page['url'],
                'original_content': original_content,
                'content': content,
                'level': level,
                'subpages': []
            }

            page_info = [info]

            if 'subPages' in page:
                for sub_page in page['subPages']:
                    page_info.extend(extract_pages_recursive(sub_page, headers, level+1))

            return page_info

        def extract_url_values(url):
            regex = r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/(?P<page_id>\d+)/.*"
            match = re.search(regex, url)
            if match:
                return match.groupdict()
            else:
                return None

        def obtain_page(organization, project, wiki, page_id, headers):
            api_url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki}/pages/{page_id}?api-version=7.0&recursionLevel=full&includeContent=true"
            response = requests.get(api_url, headers=headers)

            if response.status_code == 200:
                page = json.loads(response.text)
                return page
            else:
                print(f"Error al obtener la página de la Wiki: {response.status_code}")

            return None

        def download_specific_page(headers, wiki_url):
            url_values = extract_url_values(wiki_url)

            if url_values is None:
                print("La URL proporcionada no es válida.")
                return

            page_id = url_values['page_id']
            organization = url_values['organization']
            project = url_values['project']
            wiki = url_values['wiki']

            page = obtain_page(organization, project, wiki, page_id, headers)

            if page is not None:
                page_info = extract_pages_recursive(page, headers)
                

                # Escribir el contenido original de las páginas en un archivo Markdown
                with open('htmlymd.md', 'w', encoding='utf-8') as f:
                    for page in page_info:
                        # Obtener el título y contenido original de la página
                        title = page['name'].split('/')[-1]
                        original_content = page['original_content'].strip()

                        # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                        if original_content:
                            f.write(f'# {title}\n')
                            f.write(original_content)
                            f.write('\n\n')

                        # Recorrer las subpáginas y escribir su contenido original
                        for subpage in page['subpages']:
                            # Obtener el título y contenido original de la subpágina
                            subpage_title = subpage['name'].split('/')[-1]
                            subpage_original_content = subpage['original_content'].strip()

                            # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                            if subpage_original_content:
                                f.write(f'# {subpage_title}\n')
                                f.write(subpage_original_content)
                                f.write('\n\n')
                
                        
                # Escribir el contenido de las páginas en un archivo Markdown
                with open('todosmd.md', 'w', encoding='utf-8') as f:
                    for page in page_info:
                        # Obtener el título y contenido de la página
                        title = page['name'].split('/')[-1]
                        content = page['content'].strip()
                        # Eliminar espacios adicionales del contenido de las páginas
                        content = re.sub(r'\s+', ' ', content)

                        # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                        if content:
                            f.write(f'# {title}\n')
                            f.write(content)
                            f.write('\n\n')

                        # Recorrer las subpáginas y escribir su contenido
                        for subpage in page['subpages']:
                            # Obtener el título y contenido de la subpágina
                            subpage_title = subpage['name'].split('/')[-1]
                            subpage_content = subpage['content'].strip()
                            # Eliminar espacios adicionales del contenido de las subpáginas.
                            content = re.sub(r'\s+', ' ', content)

                            # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                            if subpage_content:
                                f.write(f'# {subpage_title}\n')
                                f.write(subpage_content)
                                f.write('\n\n')
                                
            else:
                print(f"No se encontró la página con ID {page_id}.")
            
            return page_info

        wiki_url = input("Introduce la URL de la página que quieres buscar: ")
        personal_access_token = input("Introduce tu token de acceso personal: ")
        credentials = f":{personal_access_token}"
        encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')
        headers = {
                'Authorization': f'Basic {encoded_credentials}',
                "Content-Type": "application/json",
                "Accept": "application/json"
        }
        page_info = download_specific_page(headers, wiki_url)

        def extract_placeholders(template_path):
            doc = Document(template_path)
            placeholders = []

            previous_full_text = None  # Agregar variable para almacenar el texto completo anterior

            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if full_text != previous_full_text:  # Solo imprimir el texto completo si es diferente del anterior
                            previous_full_text = full_text  # Actualizar el valor de previous_full_text
                        
                        # Find placeholders in the full text
                        start_index = full_text.find('{{')
                        end_index = full_text.find('}}')
                        while start_index != -1 and end_index != -1:
                            placeholder = full_text[start_index+2:end_index].strip()
                            placeholders.append(placeholder)
                            
                            # Find the next placeholder
                            start_index = full_text.find('{{', end_index)
                            end_index = full_text.find('}}', end_index+2)

                elif element.tag.endswith('tc'):  # Check for table cell (td)
                    for p in element.iterchildren('{%s}p' % nsmap['w']):
                        paragraph = Paragraph(p, doc)
                        if hasattr(paragraph, 'runs'):
                            # Concatenate the text of adjacent runs
                            full_text = ''.join([run.text for run in paragraph.runs])
                            if full_text != previous_full_text:  # Solo imprimir el texto completo si es diferente del anterior
                                previous_full_text = full_text  # Actualizar el valor de previous_full_text
                            
                            # Find placeholders in the full text
                            start_index = full_text.find('{{')
                            end_index = full_text.find('}}')
                            while start_index != -1 and end_index != -1:
                                placeholder = full_text[start_index+2:end_index].strip()
                                placeholders.append(placeholder)
                                
                                # Find the next placeholder
                                start_index = full_text.find('{{', end_index)
                                end_index = full_text.find('}}', end_index+2)

            # Remove duplicates from the list of placeholders while maintaining the order of the elements
            placeholders = list(OrderedDict.fromkeys(placeholders))

            return placeholders

        def create_context(page_info, placeholders):
            context = {}
            title_index = 0
            for placeholder in placeholders:
                if title_index < len(page_info):
                    page = page_info[title_index]
                    # Agregar valores al contexto tanto para marcadores de posición de título como para marcadores de posición de contenido
                    if placeholder.endswith('_content'):
                        # Eliminar espacios adicionales del contenido
                        content = re.sub(r'\s+', ' ', page['content'])
                        # Eliminar las dos almohadillas del contenido
                        content = content.replace('##', '')
                        context[placeholder] = page['content']
                        title_index += 1
                    else:
                        context[placeholder] = page['name'].split('/')[-1]
            return context


        def add_titles_to_template(template_path, titles_info):
            doc = Document(template_path)

            # Find the element after which to insert the placeholders
            insert_element = None
            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if 'REFERENCES' in full_text:
                            insert_element = paragraph
                            break

            if insert_element is not None:
                # Insert the titles before the element with the desired heading style
                for title_info in titles_info:
                    sanitized_title = sanitize_placeholder(title_info["title"])
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}}}}}')
                    # Map levels to Word Heading styles
                    level_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4', 5: 'Heading 5'}
                    heading_style = level_map.get(title_info["level"], 'Normal')
                    p.style = heading_style

                    # Insert a placeholder for the content below the title
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}_content}}}}')
                    # Set the alignment of the paragraph to left
                    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Apply bold formatting to the text if it starts with two hash symbols and remove the hash symbols from the content
                    if p.text.startswith('##'):
                        for run in p.runs:
                            run.bold = True
                            run.text = run.text.replace('##', '')

                # Insert a page break after the last title
                p = insert_element.insert_paragraph_before()
                run = p.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)

            # Save the modified template
            doc.save('new_template.docx')



        # Generate the title information
        titles_info = [{'title': page['name'].split('/')[-1], 'level': page['level']} for page in page_info if page['content'].strip() and page['name'].split('/')[-1].strip()]


        # Call the function to create a new template and add the titles to it
        add_titles_to_template('plantilla.docx', titles_info)

        # Load the modified template
        doc = DocxTemplate("new_template.docx")

        # Extract the placeholders from the template
        placeholders = extract_placeholders("new_template.docx")

        # Create a new context
        context = create_context(page_info, placeholders)


        # Render template with dynamic context
        doc.render(context)

        # Save the generated document
        doc.save("documento_generado.docx")


        print("Wait one moment to finish work")
        # Wait for a few seconds to make sure the file is saved
        time.sleep(4)
        print("Ya, todo ready by GPT")


        # Check if the file was saved successfully   
        if os.path.exists("documento_generado.docx"):
            print("El archivo documento_generado.docx se ha guardado correctamente.")
        else:
            print("No se pudo guardar el archivo documento_generado.docx.")

        # Get the directory of the script
        dir_path = os.path.dirname(os.path.realpath(__file__))

        # Build the file path
        docx_file = os.path.join(dir_path, "documento_generado.docx")

        # Call the function to update the table of contents in the generated document
        update_toc(docx_file)

        document = Document('documento_generado.docx')


    paginaconcreta2()  # Llama a la función paginaconcreta2


print("Se le solicita al usuario estos datos para acceder git azure,ficheros adjuntos.Pon los datos bien aunque no tengas datos adjuntos para que el programa siga su curso hasta finalizar.")

# Solicita al usuario el token de acceso y la URL de la wiki
wiki_url = input("Introduzca otra vez la URL principal de la wiki de Azure: ")
personal_access_token = input("Introduzca su token de acceso personal: ")


try:

   # Analiza la URL para obtener la organización y el nombre del repositorio
   parsed_url = urllib.parse.urlparse(wiki_url)
   path_parts = parsed_url.path.split("/")
   organization = path_parts[1]
   repository = path_parts[3]
except IndexError:
   print("Por favor Usuario, asegúrese de introducir correctamente la URL y el Token de Acceso.")
   exit()


# URL de la API de Git para tu proyecto y wiki
list_url = "https://dev.azure.com/ivosanchez0159/Prueba/_apis/git/repositories/Prueba.wiki/items?scopePath=/.attachments&recursionLevel=full&api-version=5.0"

# Configura la autenticación
headers = {
    'Authorization': f'Basic {base64.b64encode((":{}".format(personal_access_token)).encode()).decode()}'
}

# Crea la carpeta .attachments si no existe
if not os.path.exists('.attachments'):
    os.makedirs('.attachments')

# Realiza la solicitud GET para obtener la lista de archivos
list_response = requests.get(list_url, headers=headers)

# Verifica si el estado de la respuesta es 203 (autenticación fallida)
if list_response.status_code == 203:
    print("Por favor Usuario, asegúrese de introducir correctamente el Token de Acceso o URL.")
    exit()

# Procesa la respuesta
if list_response.status_code == 200:
    response_json = list_response.json()
    if 'value' in response_json:
        files = response_json['value']
        for file in files:
            # Salta el elemento que representa la carpeta en sí
            if file.get('isFolder'):
                continue

            # Descarga cada archivo
            download_url = f"https://dev.azure.com/ivosanchez0159/Prueba/_apis/git/repositories/Prueba.wiki/items?path={file['path']}&api-version=5.0"
            download_response = requests.get(download_url, headers=headers)
            
            if download_response.status_code == 200:
                file_name = os.path.basename(file['path'])
                with open(f'.attachments/{file_name}', 'wb') as f:
                    f.write(download_response.content)
                print(f"Descargado con éxito: {file_name}")
            else:
                print(f"Error al descargar {file['path']}: {download_response.status_code}")
    else:
        print("La clave 'value' no está presente en la respuesta de la API.")
else:
    print(f"Error al listar archivos: {list_response.status_code}")

# Listar los nombres de los archivos en el directorio .attachments
print("Archivos en el directorio .attachments:")
for filename in os.listdir('.attachments'):
    print(filename)



print("Iniciando la aplicación de Word y abriendo el documento...")

# Inicializar la aplicación de Word y abrir el documento
word_app = win32com.client.Dispatch("Word.Application")
word_app.Visible = False

file_path = os.path.abspath("documento_generado.docx")
doc = word_app.Documents.Open(file_path)

# Patrón específico que identifica la estructura que nos interesa
specific_pattern = re.compile(r'(\[.*?\]\(/.*?\))\s*(\d+\.\d+\.\s*.*?$)')

# Compilar las expresiones regulares para los patrones de enlace
link_patterns = [
    re.compile(r'\[.*\]\(/.*\)'),  # Enlace tipo1: System-Information-Leak
    re.compile(r'- \[.*\]\(/.*\)'),  # Enlace tipo2: - System-Information-Leak
    re.compile(r'- https://.*'),  # Enlace tipo3: - https://www.enlaceprueba.com
    re.compile(r'!\[.*\]\(https://.*\)')  # Enlace tipo4: !ejemplo otro tipo de enlace
]

for para in doc.Paragraphs:
    match = specific_pattern.search(para.Range.Text)
    if match:
        # Separar el enlace del título
        para.Range.Text = match.group(1) + '\n' + match.group(2)
        para.Range.Text = match.group(1) + '\n' + match.group(2)
        print(f"Insertando salto de párrafo entre '{match.group(1)}' y '{match.group(2)}'...")
    else:
        # Si no coincide con el patrón específico, busca otros patrones de enlace
        for pattern in link_patterns:
            if pattern.search(para.Range.Text):
                end_pos = para.Range.End
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo

# Una segunda pasada para eliminar párrafos vacíos que tienen numeración
for para in doc.Paragraphs:
    if para.Range.Text.strip() == '' and para.Range.ListFormat.ListType != 0:  # Párrafo vacío con numeración
        para.Range.Delete()


find_object = doc.Content.Find
find_object.ClearFormatting()
find_object.Text = '^l'  # '^l' es el código para '↵'
find_object.Replacement.ClearFormatting()
find_object.Replacement.Text = '^p'  # '^p' es el código para '¶'
find_object.Execute(Replace=2)  # 2 = wdReplaceAll




# Lista para almacenar todas las tablas encontradas
all_tables_data = []

# Patrón regex para encontrar enlaces en Markdown
pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')

# Patrón regex para encontrar imágenes en Markdown rodeadas por tuberías
image1_pattern1 = re.compile(r'\|\!\[([^\]]*)\]\((.*?)\)\|')

table = None  # Añade esta línea para inicializar 'table' como None

# Analizar cada párrafo buscando el inicio y el final de una tabla en markdown
while True:
    start_table = None
    end_table = None
    table_lines = []

    # Búsqueda de tablas en el documento
    for index, para in enumerate(doc.Paragraphs):
        line = para.Range.Text
        placeholders = []  # Añade esta línea para definir 'placeholders'
        
        # Verifica si la línea contiene una imagen en formato Markdown rodeada por tuberías
        if image1_pattern1.search(line):
            continue  # Si es así, ignora la línea y pasa a la siguiente
        
        if "|" in line:
            if start_table is None:
                start_table = index
            # Vuelve a insertar los enlaces originales
            for placeholder, match in placeholders:
                line = line.replace(placeholder, match[0])
            table_lines.append(line.strip())
        elif start_table is not None:
            end_table = index
            break

    # Si no se encuentra ninguna tabla, se detiene el bucle
    if not table_lines:
        break

    # Procesamiento de la tabla en markdown
    data = []
    headers_found = False
    for line in table_lines:
        if "|--" in line:
            headers_found = True
            continue
        if not headers_found:
            headers = [cell.strip() for cell in line.split('|')[1:]]
            if headers[-1].strip() == '':
                headers = headers[:-1]
            continue
        else:
            cells = [cell.strip() for cell in line.split('|')[1:]]
            if cells[-1].strip() == '':
                cells = cells[:-1]
            data.append(cells)

    data.insert(0, headers)

    counter = 0
    for para in doc.Paragraphs:
        counter += 1
        if counter == start_table:
            start_range = para.Range.Start
        if counter == end_table:
            end_range = para.Range.End
            break

    doc.Range(start_range, end_range).Delete()

    table_range = doc.Range(start_range, start_range)
    
    # Comprueba si hay datos antes de crear la tabla.
    if len(data) > 0 and len(data[0]) > 0:
        table = doc.Tables.Add(table_range, len(data), len(data[0]))
    else:
        print("No se encontraron datos para crear la tabla.")
        continue


    markdown_link_found_in_table = False  # Añade esta línea para inicializar 'markdown_link_found_in_table' en False

    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            cell = table.Cell(i+1, j+1)
            cell_range = cell.Range
            
            
            cell_range.Text = cell_data.strip()
            
            
            # Manipulación de hipervínculos
            matches = pattern.findall(cell_data)
            
            if matches:
                markdown_link_found_in_table = True  # Se encontró un enlace Markdown en una celda de la tabla

                for text, url in matches:
                    hyperlink_range = cell_range.Duplicate
                  
                    # Limpia el texto de la celda antes de añadir el hipervínculo.
                    cell_range.Text = text.strip()
                    
                    
                    hyperlink_range.Find.Execute(FindText=text)
                    doc.Hyperlinks.Add(Anchor=hyperlink_range, Address=url)
                    
                    

    table.Style = "Acc_Table_1"
    all_tables_data.append(data)


# Ajusta el tamaño de las celdas a 1.76 cm (convertido a puntos)
cm_to_points = 2.25 * 28.3465  # 1 cm es aproximadamente 28.3465 puntos


if table is not None:  # Comprueba si 'table' no es None
    for row in table.Rows:
        for cell in row.Cells:
            cell.Width = cm_to_points


# Inicializa 'sheet_resized' como False
sheet_resized = False

# Recorrer las tablas y aumentar el tamaño de la hoja si hay más de 5 columnas
for table in doc.Tables:
    if table.Columns.Count > 5:
        # Aumentar el tamaño de la hoja en 2.54 cm (equivalente a 1 pulgada)
        points_in_cm = 5.45 * 28.3465  # 1 cm = 28.3465 puntos
        doc.PageSetup.PageWidth = doc.PageSetup.PageWidth + points_in_cm
        doc.PageSetup.PageHeight = doc.PageSetup.PageHeight + points_in_cm
        print(f"Se ha aumentado el tamaño de la hoja. Nuevo ancho: {doc.PageSetup.PageWidth}, Nuevo alto: {doc.PageSetup.PageHeight}")
        sheet_resized = True
        break  # Salir del bucle después de la primera tabla encontrada


if sheet_resized:
    print("El tamaño de la hoja del documento se ha redimensionado.")
else:
    print("El tamaño de la hoja del documento no ha sido alterado.")


# Compilar la expresión regular para buscar imágenes en formato Markdown
image_pattern_with_pipes = re.compile(r'\|!\[([^\]]*)\]\(([^)]+)\)\|')
image_pattern_without_pipes = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')

script_dir = os.path.dirname(os.path.abspath(__file__))

# Comprobar si el directorio '.attachments' existe
attachments_dir = os.path.join(script_dir, ".attachments")
if not os.path.isdir(attachments_dir):
    print(f"El directorio de imágenes no existe: {attachments_dir}")
else:
    print(f"El directorio .attachments existe: {attachments_dir}")

    # Enumerando todos los archivos en el directorio .attachments
    attachment_files = os.listdir(attachments_dir)


try:
    # Buscar y reemplazar imágenes en formato Markdown con imágenes incrustadas
    for paragraph in doc.Paragraphs:
        match_with_pipes = image_pattern_with_pipes.search(paragraph.Range.Text)
        match_without_pipes = image_pattern_without_pipes.search(paragraph.Range.Text)
        if match_with_pipes or match_without_pipes:
            if match_with_pipes:
                description = match_with_pipes.group(1)
                image_path_markdown = match_with_pipes.group(2)
                # Aquí va el código para manejar imágenes con pipelines alrededor
            elif match_without_pipes:
                description = match_without_pipes.group(1)
                image_path_markdown = match_without_pipes.group(2)
                # Aquí va el código para manejar imágenes sin pipelines alrededor
          #    if description:  # Asegúrate de que la descripción exista
       #            desc_paragraph = doc.Paragraphs.Add(paragraph.Range)
     #              desc_paragraph.Range.Text = f"\n{description}"
    #               desc_paragraph.Format.Alignment = win32.constants.wdAlignParagraphLeft
 
            # Ignorar enlaces externos
            if image_path_markdown.startswith("http") or image_path_markdown.startswith("https"):
                continue

            # Extraer solo el nombre del archivo y la extensión del enlace de Markdown
            file_name_ext = os.path.basename(image_path_markdown)

            # Comprobar si el archivo se encuentra en la lista 'attachment_files'
            if file_name_ext in attachment_files:
                image_path = os.path.join(attachments_dir, file_name_ext)

                paragraph.Range.Delete()

                # Agregar "|" antes de la imagen si es necesario
                if match_with_pipes:
                    pipe_paragraph_before = doc.Paragraphs.Add(paragraph.Range)
                    pipe_paragraph_before.Range.Text = "|"
                    pipe_paragraph_before.Format.Alignment = win32.constants.wdAlignParagraphLeft

                image_paragraph = doc.Paragraphs.Add(paragraph.Range if not match_with_pipes else pipe_paragraph_before.Range)
                image_paragraph.Format.Style = doc.Styles.Item("Normal")

                image_range = image_paragraph.Range
                image_range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
 
                if os.path.exists(image_path):
                    print(f"La ruta de la imagen existe: {image_path}")
                    try:
                        image = image_range.InlineShapes.AddPicture(FileName=image_path, LinkToFile=False, SaveWithDocument=True)
                    except Exception as e:
                        print(f"Error al insertar la imagen: {e}")
                        continue
                else:
                    print(f"El archivo de imagen no se encuentra en la ruta especificada: {image_path}")
                    continue

                max_height = 6 * 28.3465
                if image.Height > max_height:
                    image.Height = max_height

                # Agregar "|" después de la imagen si es necesario
                if match_with_pipes:
                    pipe_paragraph_after = doc.Paragraphs.Add(image_range)
                    pipe_paragraph_after.Range.Text = "|"
                    pipe_paragraph_after.Format.Alignment = win32.constants.wdAlignParagraphLeft

            #   if description:
          #          desc_paragraph = doc.Paragraphs.Add(pipe_paragraph_after.Range if not match_with_pipes else image_range)
       #             desc_paragraph.Range.Text = f"\n{description}"
     #               desc_paragraph.Format.Alignment = win32.constants.wdAlignParagraphLeft

except Exception as e:
    print(f"Capturada una excepción al procesar imágenes: {e}")
    print(f"Excepción capturada en el párrafo: {paragraph.Range.Text}")
    print(f"Detalles de la excepción: {e}")
           


# Compilar la expresión regular para buscar imágenes en formato Markdown con URL completa
image_pattern = re.compile(r'!\[([^\]]*)\]\((http[s]?://[^)]+)\)')


# Añadir una variable al inicio de tu script para llevar un seguimiento
message_shown = False

# Buscar y reemplazar imágenes en formato Markdown con imágenes incrustadas
for paragraph in doc.Paragraphs:
    match = image_pattern.search(paragraph.Range.Text)
    if match:
        # Si el mensaje aún no se ha mostrado, mostrarlo y actualizar la variable
        if not message_shown:
            print("Por favor espere un momento usuario, está realizando el script tareas muy complejas...")
            message_shown = True

        # Obtener la descripción y la URL de la imagen
        description = match.group(1)
        image_url = match.group(2)

        paragraph.Range.Delete()

        # Crear un nuevo párrafo con estilo "Normal" para insertar la imagen
        image_paragraph = doc.Paragraphs.Add(paragraph.Range)
        image_paragraph.Format.Style = doc.Styles.Item("Normal")

        
        image_range = image_paragraph.Range
        image_range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
        image = image_range.InlineShapes.AddPicture(FileName=image_url, LinkToFile=False, SaveWithDocument=True)

        # Ajustar la altura de la imagen a un máximo de 6 cm (6 cm * 28.3465 puntos/cm = 170.078 puntos)
        max_height = 6 * 28.3465
        if image.Height > max_height:
            image.Height = max_height

        # Agregar un subtítulo a la imagen si se proporcionó una descripción
        if description:        
            # Crear un nuevo párrafo con estilo "Normal" para insertar la descripción
            desc_paragraph = doc.Paragraphs.Add(image_range)
            desc_paragraph.Range.Text = f"\n{description}"
            desc_paragraph.Format.Alignment = win32.constants.wdAlignParagraphCenter



# Variables para detección de bloques de código
found_codeblock = False
in_codeblock = False

# Iteración sobre párrafos
for paragraph in doc.Paragraphs:
    text = paragraph.Range.Text

    # Verificar si hemos encontrado el inicio de un bloque de código
    if '```' in text:
        in_codeblock = not in_codeblock
        found_codeblock = True
        
        # Eliminar el párrafo que contiene "```"
        paragraph.Range.Delete()
        continue  # Saltar el delimitador

    # Si estamos dentro de un bloque de código, cambiar la fuente y el color del texto y del fondo
    if in_codeblock:
        paragraph.Range.Font.Name = "Consolas"
        paragraph.Range.Font.Color = win32api.RGB(255, 255, 255)  # Blanco
        paragraph.Range.Shading.BackgroundPatternColor = win32api.RGB(0, 0, 0)  # Negro

    # Si estamos dentro o fuera del bloque, eliminar los caracteres ``` antes y después.
    if '´´´' in text and not in_codeblock:  
        updated_text = text.replace('´´´', '')
        paragraph.Range.Text = updated_text


# Compilamos las expresiones regulares
markdown_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
url_pattern = re.compile(r'-\s(http[s]?://\S+)')
markdown_link_regex = re.compile(r'-\s\[(.+?)\]\((http[s]?://\S+)\)')


for i in range(1, doc.Paragraphs.Count + 1):
    paragraph = doc.Paragraphs.Item(i)
    match = markdown_pattern.search(paragraph.Range.Text)
    if match:
        # Comprobar si el siguiente párrafo es un título
        next_paragraph = doc.Paragraphs.Item(i + 1) if i + 1 <= doc.Paragraphs.Count else None
        if next_paragraph and next_paragraph.Style.NameLocal.startswith("Heading"):
            pass  # Puedes reemplazar 'pass' con tu propio código
        else:
            # Agrega un salto de párrafo después del enlace de Markdown
            paragraph.Range.InsertAfter('\r')


for paragraph in doc.Paragraphs:
    match = re.search(markdown_link_regex, paragraph.Range.Text)
    if match:
        name = match.group(1)
        url = match.group(2)
        hyperlink_range = paragraph.Range
        doc.Hyperlinks.Add(hyperlink_range, url, TextToDisplay=name)
      
for paragraph in doc.Paragraphs:
    match = markdown_pattern.search(paragraph.Range.Text)
    if match:
        name = match.group(1)
        hyperlink_url = ""

        # Comprueba si el texto de anclaje del enlace de Markdown contiene una dirección web
        if re.match(r'https?://.+', match.group(2)):
            hyperlink_url = match.group(2)

        # Si el texto de anclaje del enlace de Markdown contiene una dirección web, crea un hipervínculo de Word
        if hyperlink_url:

            paragraph.Range.Delete()

            hyperlink_range = paragraph.Range
            doc.Hyperlinks.Add(hyperlink_range, hyperlink_url, TextToDisplay=name)


for paragraph in doc.Paragraphs:
    match = url_pattern.search(paragraph.Range.Text)
    if match:
        url = match.group(1)
        hyperlink_range = paragraph.Range
        # Añadir un bucle de comprobación antes de llamar al método doc.Hyperlinks.Add
        doc.Hyperlinks.Add(hyperlink_range, url)
      

for par in doc.Paragraphs:
    paragraph = par.Range.Text    
    if paragraph.startswith('- '):
        par.Format.SpaceAfter = 0
        par.Range.ListFormat.ApplyListTemplateWithLevel(
            ListTemplate=par.Range.Application.ListGalleries.Item(1).ListTemplates.Item(1),
            ContinuePreviousList=False,
            ApplyTo=win32.constants.wdListApplyToWholeList,
            DefaultListBehavior=win32.constants.wdWord10ListBehavior)
        par.Range.ListFormat.ListLevelNumber = 1  # Aquí se especifica el nivel de lista
    
    elif paragraph.startswith('  - '):
        par.Format.SpaceAfter = 0
        par.Range.ListFormat.ApplyListTemplateWithLevel(
            ListTemplate=par.Range.Application.ListGalleries.Item(1).ListTemplates.Item(1),
            ContinuePreviousList=False,
            ApplyTo=win32.constants.wdListApplyToWholeList,
            DefaultListBehavior=win32.constants.wdWord10ListBehavior)
        par.Range.ListFormat.ListLevelNumber = 2



# Definir el texto objetivo y el texto de reemplazo
target_text_1 = "###"
replacement_text_1 = ""
target_text_2 = "##"
replacement_text_2 = ""
target_text_3 = "_"
replacement_text_3 = ""
target_text_4 = "#"
replacement_text_4 = ""
target_text_5 = "*"
replacement_text_5 = ""
target_text_6 = "**"
replacement_text_6 = ""
target_text_7 = "####"
replacement_text_7 = ""



# Utilizar el objeto Find para buscar en todo el documento
find_object.ClearFormatting()

# Buscar todos los párrafos que contienen "###" y guardar su texto
found_paragraphs_1 = []
find_object.Text = target_text_1
for paragraph in doc.Paragraphs:
    if target_text_1 in paragraph.Range.Text:
        found_paragraphs_1.append(paragraph.Range.Text)

# Realizar el reemplazo de "###" por " "
if found_paragraphs_1:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_1
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_1:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_1, replacement_text_1) == paragraph.Range.Text:
            paragraph.Range.Bold = True
            paragraph.Range.Font.Size = 11
            

found_paragraphs_2 = []
find_object.Text = target_text_2
for paragraph in doc.Paragraphs:
    if target_text_2 in paragraph.Range.Text:
        found_paragraphs_2.append(paragraph.Range.Text)

# Realizar el reemplazo de "##" por " "
if found_paragraphs_2:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_2
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_2:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_2, replacement_text_2) == paragraph.Range.Text:
            paragraph.Range.Bold = True
            paragraph.Range.Font.Size = 16
           
found_paragraphs_3 = []
find_object.Text = target_text_3
for paragraph in doc.Paragraphs:
    if target_text_3 in paragraph.Range.Text:
        found_paragraphs_3.append(paragraph.Range.Text)

# Realizar el reemplazo de "_" por " "
if found_paragraphs_3:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_3
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_3:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_3, replacement_text_3) == paragraph.Range.Text:
            paragraph.Range.Italic = True
            
# Buscar todos los párrafos que contienen "#" y guardar su texto
found_paragraphs_4 = []
find_object.Text = target_text_4
for paragraph in doc.Paragraphs:
    if target_text_4 in paragraph.Range.Text:
        found_paragraphs_4.append(paragraph.Range.Text)

# Realizar el reemplazo de "#" por " "
if found_paragraphs_4:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_4
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_4:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_4, replacement_text_4) == paragraph.Range.Text:
            paragraph.Range.Bold = True
            paragraph.Range.Font.Size = 16

# Buscar todos los párrafos que contienen "**" y guardar su texto
found_paragraphs_6 = []
find_object.Text = target_text_6
for paragraph in doc.Paragraphs:
    if target_text_6 in paragraph.Range.Text:
        found_paragraphs_6.append(paragraph.Range.Text)

# Realizar el reemplazo de "**" por " "
if found_paragraphs_6:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_6
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll


for found_text in found_paragraphs_6:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_6, replacement_text_6) == paragraph.Range.Text:
            paragraph.Range.Bold = True


# Añadir el caso para "####"
found_paragraphs_7 = []
find_object.Text = target_text_7
for paragraph in doc.Paragraphs:
    if target_text_7 in paragraph.Range.Text:
        found_paragraphs_7.append(paragraph.Range.Text)

# Realizar el reemplazo de "####" por " "
if found_paragraphs_7:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_7
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_7:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_7, replacement_text_7) == paragraph.Range.Text:
            paragraph.Range.Bold = True
            paragraph.Range.Font.Size = 11

# Buscar todos los párrafos que contienen "*" y guardar su texto
found_paragraphs_5 = []
find_object.Text = target_text_5
for paragraph in doc.Paragraphs:
    if target_text_5 in paragraph.Range.Text:
        found_paragraphs_5.append(paragraph.Range.Text)

# Realizar el reemplazo de "*" por " "
if found_paragraphs_5:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_5
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll


for found_text in found_paragraphs_5:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_5, replacement_text_5) == paragraph.Range.Text:
            paragraph.Range.Italic = True

for paragraph in doc.Paragraphs:
    if "---" in paragraph.Range.Text:
        # Borrar el texto original
        paragraph.Range.Text = "\n"
        
        # Agregar una línea horizontal
        border = paragraph.Range.Borders(win32.constants.wdBorderTop)
        border.LineStyle = win32.constants.wdLineStyleSingle
        border.LineWidth = win32.constants.wdLineWidth050pt

        # Establecer el color de la línea a gris claro
        border.Color = win32api.RGB(234, 234, 234)


def apply_shading_to_text(container, pattern):
    # Iterar sobre todos los elementos en el contenedor (párrafos o celdas)
    for element in container:
        # Buscar el patrón en el texto del elemento
        matches = re.findall(pattern, element.Range.Text)
        
        # Si se encontraron coincidencias, procesar cada una
        if matches:
            for match in matches:
                # Crear un rango para la coincidencia
                start_pos = element.Range.Text.find('`' + match + '`')
                end_pos = start_pos + len(match) + 2  # Se suma 2 para incluir los caracteres de acento grave
                
                # Crear un nuevo rango que solo incluye la palabra que coincide con el patrón
                word_range = doc.Range(element.Range.Start + start_pos, element.Range.Start + end_pos)
                
                # Aplicar el sombreado gris al rango de la palabra
                word_range.Shading.BackgroundPatternColor = win32.constants.wdColorGray15
                
                # Eliminar los caracteres de acento grave
                word_range.Text = match

# Definir el patrón de búsqueda
pattern = '`(.*?)`'


apply_shading_to_text(doc.Paragraphs, pattern)

for table in doc.Tables:
    apply_shading_to_text(table.Range.Cells, pattern)


# Patrón regex para identificar URLs
url_pattern = re.compile(r'\b((http|https):\/\/)?[^\s()<>]+(?:\.[a-z]{2,})')

# Patrón regex para identificar imágenes en Markdown
image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')


# Iterar sobre todos los párrafos en el documento
for paragraph in doc.Paragraphs:
    # Buscar todas las URLs en el párrafo
    for match in re.finditer(url_pattern, paragraph.Range.Text):
        # Obtener la URL y su posición en el texto del párrafo
        url = match.group()
        start_pos = match.start()
        end_pos = match.end()

         # Comprobar si la URL es parte de una imagen en Markdown
        if image_pattern.search(paragraph.Range.Text[start_pos:end_pos]):
            print(f"Ignorando la URL '{url}' porque es parte de una descripción de imagen en Markdown.")
            continue  # Si es así, ignora la URL y pasa a la siguiente



        # Crear un rango que solo incluye la URL
        url_range = doc.Range(paragraph.Range.Start + start_pos, paragraph.Range.Start + end_pos)

        # Añadir un hipervínculo a la URL
        doc.Hyperlinks.Add(Anchor=url_range, Address=url)
        

# Itera sobre todos los campos en el documento
for field in doc.Fields:
    # Comprueba si el campo es un hipervínculo
    if field.Type == win32com.client.constants.wdFieldHyperlink:
        # Comprueba si el campo es parte de la tabla de contenido
        if field.Code.Text.startswith(" TOC "):
            # Si es parte de la tabla de contenido, no cambies el formato
            continue
        # Cambia el color del texto a azul
        field.Result.Font.Color = win32com.client.constants.wdColorBlue
        # Cambia el color del subrayado a azul
        field.Result.Font.UnderlineColor = win32com.client.constants.wdColorBlue
        # Aplica un subrayado simple
        field.Result.Font.Underline = win32com.client.constants.wdUnderlineSingle


# Recorrer todos los párrafos del documento en orden inverso
for i in range(doc.Paragraphs.Count, 0, -1):
    paragraph = doc.Paragraphs.Item(i)
    
    # Comprobar si el párrafo contiene una imagen
    if paragraph.Range.InlineShapes.Count > 0:
        # Si el párrafo anterior es un salto de párrafo, eliminarlo
        if i > 1:  # Asegurarse de que no es el primer párrafo
            prev_paragraph = doc.Paragraphs.Item(i - 1)
            if prev_paragraph.Range.Text.strip() == "":
               prev_paragraph.Range.Delete()


print("Leyendo el contenido de Markdown...")
# Leer el contenido de Markdown
with open('htmlymd.md', 'r', encoding='utf-8') as f:
    markdown_content = f.read()

def apply_formatting(doc, text_list, color_code, color_name):
    print(f"Entrando a apply_formatting con color_name = {color_name}")
    for para in doc.Paragraphs:
        run_range = para.Range
        run_text = run_range.Text
        for text in text_list:
            if text in run_text:
                start_pos = run_text.find(text)
                end_pos = start_pos + len(text)
                specific_range = doc.Range(run_range.Start + start_pos, run_range.Start + end_pos)

                # Extraer contexto alrededor del texto
                start_context = max(0, start_pos - 10)  # 10 caracteres antes
                end_context = min(len(run_text), end_pos + 10)  # 10 caracteres después
                surrounding_text = run_text[start_context:end_context]

                # Verificamos si el texto está cerca de "[Online]" y "https://"
                if re.search(r"[a-zA-Z]", run_text[start_context:start_pos]) and re.search(r"[a-zA-Z]:\/\/", run_text[end_pos:end_context]):
                    continue
                
                # Verificamos si el texto está cerca de "E-mail" y "Platform"
                if re.search(r"[a-zA-Z]", run_text[start_context:start_pos]) and re.search(r"[a-zA-Z]", run_text[end_pos:end_context]):
                    continue

                specific_range.Font.Color = color_code
                specific_range.Font.Bold = True

                


# Define los patrones de las etiquetas HTML de color
tag_patterns = {
    'red': r'<span style="color:red">(.*?)</span>',
    'green': r'<span style="color:green">(.*?)</span>',
    'teal': r'<span style="color:Teal">(.*?)</span>',
    'purple': r'<span style="color:purple">(.*?)</span>',
    'blue': r'<span style="color:blue">(.*?)</span>',
    'crimson': r'<span style="color:Crimson">(.*?)</span>',
    'bold_blue': r'<b><span style="color:blue">\s*(.*?)\s*</span></b>',
    'unclosed_red': r'<span style="color:red">\s+(.*?)(?:\s|$)',
    'unclosed2_red': r'<span style="color:red">\s*(.*?)(?=\n|$)'
}

# Diccionario con los códigos de color
color_codes = {
    'red': 255,
    'green': 65280,
    'teal': 8421376,
    'purple': 8388736,
    'blue': 16711680,
    'crimson': 139,
    'bold_blue': 16711680,
    'unclosed_red': 255,
    'unclosed2_red': 255
}

# Buscar texto y aplicar formato
for tag_name, pattern in tag_patterns.items():
    text_list = re.findall(pattern, markdown_content, re.DOTALL)
    apply_formatting(doc, text_list, color_codes[tag_name], tag_name.capitalize())


print("Limpieza de párrafos...")

# Comprobar si el documento tiene al menos 6 páginas
if doc.ComputeStatistics(win32.constants.wdStatisticPages) >= 6:
    # Obtener el rango de la sexta página
    sixth_page_range = word_app.Selection.GoTo(What=win32.constants.wdGoToPage, Which=win32.constants.wdGoToAbsolute, Count=6)
else:
    print("El documento tiene menos de 6 páginas.")
    sixth_page_range = None

# Recorrer todos los párrafos del documento en orden inverso
for i in range(doc.Paragraphs.Count, 0, -1):
    paragraph = doc.Paragraphs.Item(i)
    
    # Comprobar si el párrafo está en las primeras 5 páginas
    if sixth_page_range and paragraph.Range.Start < sixth_page_range.Start:
        continue  # Si está en las primeras 5 páginas, ignorarlo
    
    # Comprobar si el párrafo es un salto de párrafo
    if paragraph.Range.Text.strip() == "":
        # Si el párrafo anterior también es un salto de párrafo, eliminarlo
        if i > 1:  # Asegurarse de que no es el primer párrafo
            prev_paragraph = doc.Paragraphs.Item(i - 1)
            if prev_paragraph.Range.Text.strip() == "":
                try:
                    prev_paragraph.Range.Delete()
                except Exception as e:
                    print(f"No se pudo eliminar el párrafo: {e}")


# Acceder a la tabla de contenido
table_of_contents = doc.TablesOfContents(1)

# Actualizar la tabla de contenido
table_of_contents.Update()


# Guardar y cerrar el documento
doc.Save()
doc.Close()
word_app.Quit()


print("¡Finalizado!")