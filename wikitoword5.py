import os
import sys
import subprocess



# Comprobar si la carpeta 'shell' ya existe
if not os.path.exists('shell'):
    # Si no existe, crear un entorno virtual llamado "shell"
    subprocess.run([sys.executable, "-m", "venv", "shell"])

# Definir la ubicación del ejecutable de Python en el entorno virtual
venv_python = os.path.join("shell", "Scripts", "python")
if sys.platform == "linux":
    venv_python = os.path.join("shell", "bin", "python")

# Actualizar pip en el entorno virtual
subprocess.run([venv_python, "-m", "pip", "install", "--upgrade", "pip"])

# Instalar las dependencias especificadas en el archivo requirements.txt
subprocess.run([venv_python, "-m", "pip", "install", "-r", "requirements.txt"])

# Crear el segundo script que se ejecutará dentro del entorno virtual
with open("second_script.py", "w", encoding='utf-8') as f:
    f.write("""
# -*- coding: utf-8 -*-

import json
import base64
import win32com.client
import win32com.client as win32
import re
import win32api
import urllib.parse  # Importado para analizar la URL
import requests
import os
from docxtpl import DocxTemplate
from docx import Document
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
import docx
import time
from collections import OrderedDict
from win32com.client import constants
import shutil


# Pregunta al usuario si quiere descargar todas las páginas y subpáginas de la Wiki
respuesta = input("¿Deseas descargar todas las páginas de la Wiki Azure o solo una página? Por favor, responde con ‘sí’ para descargar todas las páginas, o ‘no’ para descargar solo una página? (si/no): ")

if respuesta.lower() == 'si':
    def todaslaspaginas():
        import json
        import base64
        import win32com.client
        from win32com.client import constants
        import win32com.client as win32
        import re
        import win32api
        import urllib.parse  # Importado para analizar la URL
        import requests
        import os
        from docxtpl import DocxTemplate
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsmap
        import docx
        import time
        from collections import OrderedDict



        # Inicializa las variables globales
        stored_wiki_url = None
        stored_personal_access_token = None

        def sanitize_placeholder(placeholder):
            # Replace spaces and invalid characters with underscores
            sanitized = re.sub(r'[^\w]', '_', placeholder)
            # Remove leading digits and underscores to ensure a valid variable name
            sanitized = re.sub(r'^\d+|_', '', sanitized)
            return sanitized


        def update_toc(docx_file):
            try:
                word = win32com.client.DispatchEx("Word.Application")
                doc = word.Documents.Open(docx_file)
                doc.TablesOfContents(1).Update()
                doc.Close(SaveChanges=True)
                word.Quit()
            except Exception as e:
                print(f"An error occurred while updating the table of contents: {e}")


        def get_page_content(url):
            content_url = url + "?api-version=7.0&includeContent=true"
            response = requests.get(content_url, headers=headers)

            if response.status_code == 200:
                return json.loads(response.text)['content']
            else:
                print(f"Error al obtener el contenido de la página: {response.status_code}")
                return "Contenido no disponible"

        def extract_pages_recursive(page, level=1):
            if not page:
                return []

            original_content = get_page_content(page['url'])
            content = get_page_content(page['url'])
            # Replace <span> tags with plain text
            content = re.sub(r'(<span style="color:([^>]*)">([^<]*?))\\n', r'\\1(/span)\\n', content)
            # etiqueta html completa.
            content = re.sub(r'<b><span style="color:([^>]*)">([^<]*)</span></b>', r'(b)(span style="color:\\1")\\2(/span)(/b)', content, flags=re.IGNORECASE)
            # Reemplazar <center>(.*?)</center> con '(.*?)'
            content = re.sub(r'<center>(.*?)</center>', r'\\1', content)
            # reemplazar ciertos bloques códigos wiki
            content = re.sub(r'<code>```[^\\n]*\\n', '```\\n', content)
            # Reemplazar <br>(.*?)</br> con '(.*?)'
            content = re.sub(r'<br>(.*?)</br>', r'\\1', content)
            # reemplaza etiqueta <br> con un salto de linea
            content = re.sub(r'<br>', '\\n', content)
            # Reemplazar To_do @<lo que sea> con TO_DO @lo que sea
            content = re.sub(r'TO_DO: @<([A-Fa-f0-9-]+)>', r'TO_DO: \\1', content)
            # Reemplazar <Lista> @<lo que sea> con Lista @lo que sea
            content = re.sub(r'<Lista> @<([^>]+)>', r'Lista @\\1', content)
            # quitar mayusculas formatos de imagen
            content = re.sub(r'\\.(PNG|JPG|JPEG|GIF)', lambda x: x.group().lower(), content)
            # ajustar de tal forma <hola-quepasa> se vería así (hola-quepasa)
            content = re.sub(r'<([^>]*)>', r'(\\1)', content)
           
            
            info = {
                'name': page['path'],
                'short_name': page['path'].split('/')[-1],
                'url': page['url'],
                'original_content': original_content,
                'content': content ,
                'level': level,
                'subpages': []
            }


            page_info = [info]

            if 'subPages' in page:
                for sub_page in page['subPages']:
                    page_info.extend(extract_pages_recursive(sub_page, level+1))

            return page_info

        def extract_url_values(url):
            regex = r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/.*"
            match = re.search(regex, url)
            if match:
                return match.groupdict()
            else:
                return None


        def extract_placeholders(template_path):
            doc = Document(template_path)
            placeholders = []

            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        
                        # Find placeholders in the full text
                        start_index = full_text.find('{{')
                        end_index = full_text.find('}}')
                        while start_index != -1 and end_index != -1:
                            placeholder = full_text[start_index+2:end_index].strip()
                            placeholders.append(placeholder)
                            
                            # Find the next placeholder
                            start_index = full_text.find('{{', end_index)
                            end_index = full_text.find('}}', end_index+2)

                elif element.tag.endswith('tc'):  # Check for table cell (td)
                    for p in element.iterchildren('{%s}p' % nsmap['w']):
                        paragraph = Paragraph(p, doc)
                        if hasattr(paragraph, 'runs'):
                            # Concatenate the text of adjacent runs
                            full_text = ''.join([run.text for run in paragraph.runs])
                            
                            # Find placeholders in the full text
                            start_index = full_text.find('{{')
                            end_index = full_text.find('}}')
                            while start_index != -1 and end_index != -1:
                                placeholder = full_text[start_index+2:end_index].strip()
                                placeholders.append(placeholder)
                                
                                # Find the next placeholder
                                start_index = full_text.find('{{', end_index)
                                end_index = full_text.find('}}', end_index+2)

            # Remove duplicates from the list of placeholders while maintaining the order of the elements
            placeholders = list(OrderedDict.fromkeys(placeholders))

            return placeholders


        def create_context(page_info, placeholders):
            context = {}
            title_index = 1
            for placeholder in placeholders:
                if title_index < len(page_info):
                    page = page_info[title_index]
                    # Agregar valores al contexto tanto para marcadores de posición de título como para marcadores de posición de contenido
                    if placeholder.endswith('_content'):
                        if page['content'] != 'No hay contenido':
                            context[placeholder] = page['content']
                        title_index += 1
                    else:
                        context[placeholder] = page['short_name']
            return context

        # Pedir la información al usuario una sola vez y almacenarla en variables globales
        # Verificar si las credenciales ya se han almacenado
        if stored_wiki_url is None or stored_personal_access_token is None:
            wiki_url = input("Introduce la URL principal del portal Wiki de Azure: ")
            personal_access_token = input("Introduce tu token de acceso personal: ")
            # Aquí iría la lógica para verificar si las credenciales son válidas.
            # Si son válidas, las almacenamos en las variables globales.
            stored_wiki_url = wiki_url
            stored_personal_access_token = personal_access_token
        else:
            # Usar las credenciales almacenadas
            wiki_url = stored_wiki_url
            personal_access_token = stored_personal_access_token


        credentials = f":{personal_access_token}"
        encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')

        headers = {
            'Authorization': f'Basic {encoded_credentials}',
            "Content-Type": "application/json",
            "Accept": "application/json"
        }

        # Decodifica la URL para manejar caracteres especiales
        decoded_wiki_url = urllib.parse.unquote(wiki_url)
        # Usa regex para extraer la organización, el proyecto y el wiki de la URL decodificada
        url_match = re.search(r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/.*", decoded_wiki_url)
        if url_match:
            organization = url_match.group('organization')
            project = url_match.group('project')
            wiki = url_match.group('wiki')
        else:
            print("The URL provided is not valid.")
            exit()

        url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki}/pages?api-version=7.0&recursionLevel=full"

        response = requests.get(url, headers=headers)

        if response.status_code == 200:
            print("Successful response http 200.")
            root_page = json.loads(response.text)
        else:
            print(f"Error al obtener la página raíz de la wiki: {response.status_code}")
            root_page = {}

        page_info = extract_pages_recursive(root_page)


        # Save the .md file with original content
        md_filename_original = 'htmlymd.md'
        with open(md_filename_original, 'w', encoding='utf-8') as f:
            for page in page_info:
                # Get the title and original content of the page
                title = page['name'].split('/')[-1]
                original_content = page['original_content'].strip()

                # Only write the title and original content to the Markdown file if the original content is not empty
                if original_content and not (title.startswith("#") and "No hay contenido" in title) and page['name'] != "/":
                    f.write(f'# {title}\\n')
                    f.write(original_content)
                    f.write('\\n\\n')


        # Save the .md file
        md_filename = 'todosmd.md'
        with open('todosmd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Get the title and content of the page
                title = page['name'].split('/')[-1]
                content = page['content'].strip()

                # Only write the title and content to the Markdown file if the content is not empty
                if content and not (title.startswith("#") and "No hay contenido" in title) and page['name'] != "/":
                    f.write(f'# {title}\\n')
                    f.write(content)
                    f.write('\\n\\n')


        # Read the Markdown content
        with open(md_filename, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        # Extract the level 1 headings with regex
        title_pattern = r'^# (.*)'
        titles = re.findall(title_pattern, markdown_content, flags=re.M)


        # Generate the placeholders
        placeholders = [{'text': f'Titulo{i+1}', 'level': page['level']} for i, page in enumerate(page_info)]

        def add_titles_to_template(template_path, titles_info):
            doc = Document(template_path)

            # Find the element after which to insert the placeholders
            insert_element = None
            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if 'REFERENCES' in full_text:
                            insert_element = paragraph
                            break

            if insert_element is not None:
                # Insert the titles before the element with the desired heading style
                for title_info in titles_info:
                    # Skip titles with name '/' or empty short_name
                    if 'name' in title_info and (title_info['name'] == '/' or not title_info['short_name']):
                        continue
                    sanitized_title = sanitize_placeholder(title_info["title"])
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}}}}}')
                    # Calculate the maximum level in the JSON
                    max_level = max([title_info['level'] for title_info in titles_info])
                    # Generate the level_map dictionary dynamically
                    level_map = {i: f'Heading {i}' for i in range(1, max_level+1)}
                    heading_style = level_map.get(title_info["level"] - 1, 'Normal')
                    p.style = heading_style
                    # Insert a placeholder for the content below the title
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}_content}}}}')

                # Insert a page break after the last title
                p = insert_element.insert_paragraph_before()
                run = p.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)

            # Save the modified template
            doc.save('new_template.docx')


        # Generate the title information
        titles_info = [{'title': page['name'].split('/')[-1], 'level': page['level']} for page in page_info if page['content'].strip() and page['name'].split('/')[-1].strip()]


        # Call the function to create a new template and add the titles to it
        add_titles_to_template('plantilla.docx', titles_info)

        # Load the modified template
        doc = DocxTemplate("new_template.docx")

        # Extract the placeholders from the template
        placeholders = extract_placeholders("new_template.docx")


        # Create a new context
        context = create_context(page_info, placeholders)


        # Render template with dynamic context
        doc.render(context)

        # Save the generated document
        doc.save("documento_generado.docx")


        print("Wait one moment to finish work")
        # Wait for a few seconds to make sure the file is saved
        time.sleep(4)
        print(" Already ")

        # Check if the file was saved successfully
        if os.path.exists("documento_generado.docx"):
            print("The generated_document.docx file has been saved successfully.")
        else:
            print("No se pudo guardar el archivo documento_generado.docx.")

        # Get the directory of the script
        dir_path = os.path.dirname(os.path.realpath(__file__))

        # Build the file path
        docx_file = os.path.join(dir_path, "documento_generado.docx")

        # Call the function to update the table of contents in the generated document
        update_toc(docx_file)  
       

        list_url = f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{wiki}/items?scopePath=/.attachments&recursionLevel=full&api-version=5.0"

        # Configura la autenticación
        headers = {
            'Authorization': f'Basic {base64.b64encode((":{}".format(personal_access_token)).encode()).decode()}'
        }

        # Crea la carpeta .attachments si no existe
        if not os.path.exists('.attachments'):
            os.makedirs('.attachments')

        # Realiza la solicitud GET para obtener la lista de archivos
        list_response = requests.get(list_url, headers=headers)

        # Procesa la respuesta
        if list_response.status_code == 200:
            response_json = list_response.json()
            if 'value' in response_json:
                files = response_json['value']
                folders = [file for file in files if file.get('isFolder')]
                for folder in folders:
                    print(f"Your folder is {folder['path']}")
                for file in files:
                    # Salta el elemento que representa la carpeta en sí
                    if file.get('isFolder'):
                        continue

                    # Descarga cada archivo
                    download_url = f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{wiki}/items?path={file['path']}&api-version=5.0"
                    download_response = requests.get(download_url, headers=headers)
                    if download_response.status_code == 200:
                        file_name = os.path.basename(file['path'])
                        with open(f'.attachments/{file_name}', 'wb') as f:
                            f.write(download_response.content)
                        print(f"Successfully downloaded the file {file_name} de la carpeta {folder['path']}")
                    else:
                        print(f"Error al descargar {file['path']} de la carpeta {folder['path']}: {download_response.status_code}")
            else:
                print("La clave 'value' no está presente en la respuesta de la API.")
        else:
            print(f"Error listing files: {list_response.status_code}")

        pass
        
    todaslaspaginas()  # Llama a la función todaslaspaginas

else:

    def paginaconcreta2():
        import requests
        import json
        import base64
        import re
        from docxtpl import DocxTemplate
        import os
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsmap
        from collections import OrderedDict
        import docx
        import win32com.client
        import time
        import urllib.parse

       
        # Inicializa las variables globales
        stored_wiki_url = None
        stored_personal_access_token = None

        def sanitize_placeholder(placeholder):
            # Replace spaces and invalid characters with underscores
            sanitized = re.sub(r'[^\w]', '_', placeholder)
            # Remove leading digits and underscores to ensure a valid variable name
            sanitized = re.sub(r'^\d+|_', '', sanitized)
            return sanitized

        def update_toc(docx_file):
            word = win32com.client.DispatchEx("Word.Application")
            doc = word.Documents.Open(docx_file)
            doc.TablesOfContents(1).Update()
            doc.Close(SaveChanges=True)
            word.Quit()


        def get_page_content(url, headers):
            content_url = url + "?api-version=7.0&includeContent=true"
            response = requests.get(content_url, headers=headers)

            if response.status_code == 200:
                # Convierte la respuesta en un objeto JSON
                response_json = json.loads(response.text)
                # Devuelve el valor del campo 'content'
                return response_json['content']
            else:
                print(f"Error al obtener el contenido de la página: {response.status_code}")
                return "Contenido no disponible"

        def extract_pages_recursive(page, headers=None, level=1):
            if not page:
                return []

            original_content = get_page_content(page['url'], headers)
            content = get_page_content(page['url'], headers)
            # Replace <span> tags with plain text
            content = re.sub(r'(<span style="color:([^>]*)">([^<]*?))\\n', r'\\1(/span)\\n', content)
            # etiqueta html completa.
            content = re.sub(r'<b><span style="color:([^>]*)">([^<]*)</span></b>', r'(b)(span style="color:\\1")\\2(/span)(/b)', content, flags=re.IGNORECASE)
            # Reemplazar <center>(.*?)</center> con '(.*?)'
            content = re.sub(r'<center>(.*?)</center>', r'\\1', content)
            # reemplazar ciertos bloques códigos wiki
            content = re.sub(r'<code>```[^\\n]*\\n', '```\\n', content)
            # Reemplazar <br>(.*?)</br> con '(.*?)'
            content = re.sub(r'<br>(.*?)</br>', r'\\1', content)
            # reemplaza etiqueta <br> con un salto de linea
            content = re.sub(r'<br>', '\\n', content)
            # Reemplazar To_do @<lo que sea> con TO_DO @lo que sea
            content = re.sub(r'TO_DO: @<([A-Fa-f0-9-]+)>', r'TO_DO: \\1', content)
            # Reemplazar <Lista> @<lo que sea> con Lista @lo que sea
            content = re.sub(r'<Lista> @<([^>]+)>', r'Lista @\\1', content)
            # quitar mayusculas formatos de imagen
            content = re.sub(r'\\.(PNG|JPG|JPEG|GIF)', lambda x: x.group().lower(), content)
            # ajustar de tal forma <hola-quepasa> se vería así (hola-quepasa)
            content = re.sub(r'<([^>]*)>', r'(\\1)', content)
            
                     
                   
            info = {
                'name': page['path'],
                'short_name': page['path'].split('/')[-1],
                'url': page['url'],
                'original_content': original_content,
                'content': content,
                'level': level,
                'subpages': []
            }

            page_info = [info]

            if 'subPages' in page:
                for sub_page in page['subPages']:
                    page_info.extend(extract_pages_recursive(sub_page, headers, level+1))

            return page_info

        def extract_url_values(url):
            regex = r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/(?P<page_id>\d+)/.*"
            match = re.search(regex, url)
            if match:
                return match.groupdict()
            else:
                return None

      
        def obtain_page(organization, project, wiki, page_id, headers):
            api_url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki}/pages/{page_id}?api-version=7.0&recursionLevel=full&includeContent=true"
            response = requests.get(api_url, headers=headers)
            if response.status_code == 200:
                print("Successful response http 200.")
                page = json.loads(response.text)
                return page
            else:
                print(f"Error al obtener la página de la Wiki: {response.status_code}")

            return None
    

        def download_specific_page(headers, wiki_url2):
            url_values = extract_url_values(wiki_url2)

            if url_values is None:
                print("La URL proporcionada no es válida.")
                return

            page_id = url_values['page_id']
            organization = url_values['organization']
            project = url_values['project']
            wiki = url_values['wiki']

            page = obtain_page(organization, project, wiki, page_id, headers)

            # Añade una condición de parada para la recursión
            if page is not None:
                return organization, project, wiki, page
            else:
                print("No se pudo descargar la página.")
                return None

              
        # Pedir la información al usuario una sola vez y almacenarla en variables globales
        # Verificar si las credenciales ya se han almacenado
        if stored_wiki_url is None or stored_personal_access_token is None or stored_wiki_url2 is None:
            wiki_url2 = input("Introduce la URL de la página que quieres buscar: ")
            personal_access_token = input("Introduce tu token de acceso personal: ")
            print("Por favor usuario necesito que especifiques tu URL principal azure wiki para poder descargar los ficheros tanto fotos u otros ficheros")
            wiki_url = input("Introduce la URL principal del portal Wiki de Azure: ")
            # Aquí iría la lógica para verificar si las credenciales son válidas.
            # Si son válidas, las almacenamos en las variables globales
            stored_wiki_url = wiki_url
            stored_personal_access_token = personal_access_token
        else:
            # Usar las credenciales almacenadas
            wiki_url = stored_wiki_url
            personal_access_token = stored_personal_access_token

        
        credentials = f":{personal_access_token}"
        encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')
        headers = {
                'Authorization': f'Basic {encoded_credentials}',
                "Content-Type": "application/json",
                "Accept": "application/json"
        }

        result = download_specific_page(headers, wiki_url2)
        if result is not None:
              organization, project, wiki, page_data = result
              page_info = extract_pages_recursive(page_data, headers)  # Asegúrate de que esto devuelva una lista 
        else:
            print("No se pudo descargar la página.")
            

        def extract_url_values(wiki_url2):
            # Decodifica la URL para manejar caracteres especiales
            decoded_wiki_url2 = urllib.parse.unquote(wiki_url2)
            # Usa regex para extraer la organización, el proyecto y el wiki de la URL decodificada
            api_url_match = re.search(r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>.+).wiki/(?P<page_id>\d+)/.*", decoded_wiki_url2)
            if api_url_match:
                 organization = api_url_match.group('organization')
                 project = api_url_match.group('project')
                 wiki = api_url_match.group('wiki')
                 page_id = api_url_match.group('page_id')
                 return {'organization': organization, 'project': project, 'wiki': wiki, 'page_id': page_id}
            else:
                 print("La URL proporcionada no es válida.")
                 exit() 

        
        # Escribir el contenido original de las páginas en un archivo Markdown
        with open('htmlymd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Obtener el título y contenido original de la página
                title = page['name'].split('/')[-1]
                original_content = page['original_content'].strip()

                # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                if original_content:
                    f.write(f'# {title}\\n')
                    f.write(original_content)
                    f.write('\\n\\n')

                # Recorrer las subpáginas y escribir su contenido original
                for subpage in page['subpages']:
                    # Obtener el título y contenido original de la subpágina
                    subpage_title = subpage['name'].split('/')[-1]
                    subpage_original_content = subpage['original_content'].strip()

                    # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                    if subpage_original_content:
                        f.write(f'# {subpage_title}\\n')
                        f.write(subpage_original_content)
                        f.write('\\n\\n')
                
        # Escribir el contenido de las páginas en un archivo Markdown
        with open('todosmd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Obtener el título y contenido de la página
                title = page['name'].split('/')[-1]
                content = page['content'].strip()

                # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                if content: 
                    f.write(f'# {title}\\n')
                    f.write(content)
                    f.write('\\n\\n')

                # Recorrer las subpáginas y escribir su contenido
                for subpage in page['subpages']:
                    # Obtener el título y contenido de la subpágina
                    subpage_title = subpage['name'].split('/')[-1]
                    subpage_content = subpage['content'].strip()

                    # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                    if subpage_content:
                        f.write(f'# {subpage_title}\\n')
                        f.write(subpage_content)
                        f.write('\\n\\n')
                        
                else:
                    if not page['content'].strip():
                            print(f"No se encontró la página con ID {page_id}.")


        def extract_placeholders(template_path):
            doc = Document(template_path)
            placeholders = []

            previous_full_text = None  # Agregar variable para almacenar el texto completo anterior

            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if full_text != previous_full_text:  # Solo imprimir el texto completo si es diferente del anterior
                            previous_full_text = full_text  # Actualizar el valor de previous_full_text
                        
                        # Find placeholders in the full text
                        start_index = full_text.find('{{')
                        end_index = full_text.find('}}')
                        while start_index != -1 and end_index != -1:
                            placeholder = full_text[start_index+2:end_index].strip()
                            placeholders.append(placeholder)
                            
                            # Find the next placeholder
                            start_index = full_text.find('{{', end_index)
                            end_index = full_text.find('}}', end_index+2)

                elif element.tag.endswith('tc'):  # Check for table cell (td)
                    for p in element.iterchildren('{%s}p' % nsmap['w']):
                        paragraph = Paragraph(p, doc)
                        if hasattr(paragraph, 'runs'):
                            # Concatenate the text of adjacent runs
                            full_text = ''.join([run.text for run in paragraph.runs])
                            if full_text != previous_full_text:  # Solo imprimir el texto completo si es diferente del anterior
                                previous_full_text = full_text  # Actualizar el valor de previous_full_text
                            
                            # Find placeholders in the full text
                            start_index = full_text.find('{{')
                            end_index = full_text.find('}}')
                            while start_index != -1 and end_index != -1:
                                placeholder = full_text[start_index+2:end_index].strip()
                                placeholders.append(placeholder)
                                
                                # Find the next placeholder
                                start_index = full_text.find('{{', end_index)
                                end_index = full_text.find('}}', end_index+2)

            # Remove duplicates from the list of placeholders while maintaining the order of the elements
            placeholders = list(OrderedDict.fromkeys(placeholders))

            return placeholders

        def create_context(page_info, placeholders):
            context = {}
            title_index = 0
            for placeholder in placeholders:
                if title_index < len(page_info):
                    page = page_info[title_index]
                    # Agregar valores al contexto tanto para marcadores de posición de título como para marcadores de posición de contenido
                    if placeholder.endswith('_content'):
                        # Eliminar espacios adicionales del contenido
                        content = re.sub(r'\s+', ' ', page['content'])
                        # Eliminar las dos almohadillas del contenido
                        content = content.replace('##', '')
                        context[placeholder] = page['content']
                        title_index += 1
                    else:
                        context[placeholder] = page['name'].split('/')[-1]
            return context


        def add_titles_to_template(template_path, titles_info):
            doc = Document(template_path)

            # Find the element after which to insert the placeholders
            insert_element = None
            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if 'REFERENCES' in full_text:
                            insert_element = paragraph
                            break

            if insert_element is not None:
                # Insert the titles before the element with the desired heading style
                for title_info in titles_info:
                    sanitized_title = sanitize_placeholder(title_info["title"])
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}}}}}')
                    # Map levels to Word Heading styles
                    level_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4', 5: 'Heading 5'}
                    heading_style = level_map.get(title_info["level"], 'Normal')
                    p.style = heading_style

                    # Insert a placeholder for the content below the title
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}_content}}}}')
                    # Set the alignment of the paragraph to left
                    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Apply bold formatting to the text if it starts with two hash symbols and remove the hash symbols from the content
                    if p.text.startswith('##'):
                        for run in p.runs:
                            run.bold = True
                            run.text = run.text.replace('##', '')

                # Insert a page break after the last title
                p = insert_element.insert_paragraph_before()
                run = p.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)

            # Save the modified template
            doc.save('new_template.docx')


        # Generate the title information
        titles_info = [{'title': page['name'].split('/')[-1], 'level': page['level']} for page in page_info if page['content'].strip() and page['name'].split('/')[-1].strip()]


        # Call the function to create a new template and add the titles to it
        add_titles_to_template('plantilla.docx', titles_info)

        # Load the modified template
        doc = DocxTemplate("new_template.docx")

        # Extract the placeholders from the template
        placeholders = extract_placeholders("new_template.docx")

        # Create a new context
        context = create_context(page_info, placeholders)

        # Render template with dynamic context
        doc.render(context)

        # Save the generated document
        doc.save("documento_generado.docx")


        print("Wait one moment to finish work")
        # Wait for a few seconds to make sure the file is saved
        time.sleep(4)
        print("already")


        # Check if the file was saved successfully   
        if os.path.exists("documento_generado.docx"):
            print("The documento_generado.docx file has been saved successfully.")
        else:
            print("No se pudo guardar el archivo documento_generado.docx.")

        # Get the directory of the script
        dir_path = os.path.dirname(os.path.realpath(__file__))

        # Build the file path
        docx_file = os.path.join(dir_path, "documento_generado.docx")

        # Call the function to update the table of contents in the generated document
        update_toc(docx_file)

        document = Document('documento_generado.docx')


        list_url = f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{wiki}/items?scopePath=/.attachments&recursionLevel=full&api-version=5.0"

        # Configura la autenticación
        headers = {
            'Authorization': f'Basic {base64.b64encode((":{}".format(personal_access_token)).encode()).decode()}'
        }

        # Crea la carpeta .attachments si no existe
        if not os.path.exists('.attachments'):
            os.makedirs('.attachments')

        # Realiza la solicitud GET para obtener la lista de archivos
        list_response = requests.get(list_url, headers=headers)

        # Procesa la respuesta
        if list_response.status_code == 200:
            response_json = list_response.json()
            if 'value' in response_json:
                files = response_json['value']
                folders = [file for file in files if file.get('isFolder')]
                for folder in folders:
                    print(f"Your folder is {folder['path']}")
                for file in files:
                    # Salta el elemento que representa la carpeta en sí
                    if file.get('isFolder'):
                        continue

                    # Descarga cada archivo
                    download_url = f"https://dev.azure.com/{organization}/{project}/_apis/git/repositories/{wiki}/items?path={file['path']}&api-version=5.0"
                    download_response = requests.get(download_url, headers=headers)
                    if download_response.status_code == 200:
                        file_name = os.path.basename(file['path'])
                        with open(f'.attachments/{file_name}', 'wb') as f:
                            f.write(download_response.content)
                        print(f"Successfully downloaded the file {file_name} de la carpeta {folder['path']}")
                    else:
                        print(f"Error al descargar {file['path']} from the folder {folder['path']}: {download_response.status_code}")
            else:
                print("La clave 'value' no está presente en la respuesta de la API.")
        else:
            print(f"Error al listar archivos: {list_response.status_code}")


    paginaconcreta2()  # Llama a la función paginaconcreta2


print("Starting the Word application and opening the document...")

try:
    # Inicializar la aplicación de Word y abrir el documento
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.DisplayAlerts = False
    word_app.Visible = False
except Exception as e:
    print(f"Se produjo un error al intentar iniciar la aplicación de Word: {e}")

try:
    file_path = os.path.abspath("documento_generado.docx")
    doc = word_app.Documents.Open(file_path)
except Exception as e:
    print(f"Se produjo un error al intentar abrir el documento: {e}")


# Patrón específico que identifica la estructura que nos interesa
specific_pattern = re.compile(r'(\[.*?\]\(/.*?\))\s*(\d+\.\d+\.\s*.*?$)')

# Compilar las expresiones regulares para los patrones de enlace
link_patterns = [
    re.compile(r'\[.*\]\(/.*\)'),  # Enlace tipo1: System-Information-Leak
    re.compile(r'- \[.*\]\(/.*\)'),  # Enlace tipo2: - System-Information-Leak
    re.compile(r'- https://.*'),  # Enlace tipo3: - https://www.enlaceprueba.com
    re.compile(r'!\[.*\]\(https://.*\)')  # Enlace tipo4: !ejemplo otro tipo de enlace
]

for para in doc.Paragraphs:
    match = specific_pattern.search(para.Range.Text)
    if match:
        # Separar el enlace del título
        para.Range.Text = match.group(1) + '\\n' + match.group(2)
        para.Range.Text = match.group(1) + '\\n' + match.group(2)
        print(f"Insertando salto de párrafo entre '{match.group(1)}' y '{match.group(2)}'...")
    else:
        # Si no coincide con el patrón específico, busca otros patrones de enlace
        for pattern in link_patterns:
            if pattern.search(para.Range.Text):
                end_pos = para.Range.End
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo
                doc.Range(end_pos, end_pos).InsertParagraphAfter()  # Inserta un salto de párrafo

# Una segunda pasada para eliminar párrafos vacíos que tienen numeración
for para in doc.Paragraphs:
    if para.Range.Text.strip() == '' and para.Range.ListFormat.ListType != 0:  # Párrafo vacío con numeración
        para.Range.Delete()

find_object = doc.Content.Find
find_object.ClearFormatting()
find_object.Text = '^l'  # '^l' es el código para '↵'
find_object.Replacement.ClearFormatting()
find_object.Replacement.Text = '^p'  # '^p' es el código para '¶'
find_object.Execute(Replace=2)  # 2 = wdReplaceAll


def verificar_etiquetas(doc):
    # Pila para rastrear las etiquetas de apertura
    pila_etiquetas = []

    for par in doc.Paragraphs:
        paragraph = par.Range.Text.strip()

        # Si encontramos una etiqueta de apertura, la añadimos a la pila
        if '(ul)' in paragraph or '(li)' in paragraph:
            pila_etiquetas.append(paragraph)

        # Si encontramos una etiqueta de cierre, comprobamos si su correspondiente etiqueta de apertura está en la parte superior de la pila
        elif '(/ul)' in paragraph or '(/li)' in paragraph:
            if pila_etiquetas and pila_etiquetas[-1] == paragraph.replace('/', ''):
                # Si la etiqueta de apertura correspondiente está en la parte superior de la pila, la eliminamos de la pila
                pila_etiquetas.pop()
            else:
                # Si la etiqueta de apertura correspondiente no está en la parte superior de la pila, tenemos una etiqueta suelta
                return True  # Devolver True si se encontró una etiqueta suelta

    # Al final, cualquier etiqueta de apertura que quede en la pila es una etiqueta suelta
    return len(pila_etiquetas) > 0  # Devolver True si se encontraron etiquetas sueltas, False en caso contrario

def eliminar_etiquetas_sueltas(doc):
    # Variables para rastrear etiquetas sueltas
    etiquetas_sueltas = []

    for par in doc.Paragraphs:
        paragraph = par.Range.Text.strip()

        # Agregar etiquetas sueltas a la lista para su posterior eliminación
        if ('(ul)' in paragraph and '(/ul)' not in paragraph) or \
           ('(li)' in paragraph and '(/li)' not in paragraph):
            etiquetas_sueltas.append(par)

    # Comprobar si se encontraron etiquetas sueltas
    if etiquetas_sueltas:
        # Eliminar párrafos con etiquetas sueltas
        for par in etiquetas_sueltas:
            par.Range.Delete()
        return True  # Devolver True si se encontraron etiquetas sueltas

    return False  # Devolver False si no se encontraron etiquetas sueltas

# Proceso principal
print("Verificando etiquetas HTML...")
if verificar_etiquetas(doc):
    print("Se encontraron etiquetas HTML sueltas. Procediendo a eliminarlas...")
    if eliminar_etiquetas_sueltas(doc):
        print("Se eliminaron etiquetas HTML sueltas. Por favor revise en su wiki el orden de las etiquetas HTML emparejadas.")
else:
    print("No se encontraron etiquetas HTML sueltas.")


print("Bullet and sub ul li HTML robot etiquetas solamente emparejadas")
# Variables para rastrear el estado de las listas, sublistas y bloques de código Markdown
en_lista = False
en_sublista = False
en_bloque_codigo = False

# Primero, verificamos si hay etiquetas sueltas
etiquetas_ul_suelta = '(ul)' in doc.Range().Text and '(/ul)' not in doc.Range().Text
etiquetas_li_suelta = '(li)' in doc.Range().Text and '(/li)' not in doc.Range().Text

for par in doc.Paragraphs:
    paragraph = par.Range.Text

    # Verificar si estamos en un bloque de código Markdown
    if '```' in paragraph:
        en_bloque_codigo = not en_bloque_codigo
        continue

    # Omitir la aplicación de formatos dentro de bloques de código
    if en_bloque_codigo:
        continue

    # Manejar etiquetas sueltas antes de procesar las listas y sublistas
    if etiquetas_ul_suelta and '(ul)' in paragraph:
        par.Range.Delete()
        continue
    if etiquetas_li_suelta and '(li)' in paragraph:
        par.Range.Delete()
        continue

    # Verificar y actualizar el estado de las listas y sublistas
    if '(ul)' in paragraph:
        en_lista = True
        par.Range.Delete()
        continue
    elif '(/ul)' in paragraph:
        en_lista = False
        par.Range.Delete()
        continue
    if '(li)' in paragraph:
        en_sublista = True
        par.Range.Delete()
        continue
    elif '(/li)' in paragraph:
        en_sublista = False
        par.Range.Delete()
        continue

    # Aplicar formato de lista o sublista según corresponda, excepto en bloques de código
    if en_lista or paragraph.startswith('- '):
        par.Format.SpaceAfter = 0
        par.Range.ListFormat.ApplyListTemplateWithLevel(
            ListTemplate=par.Range.Application.ListGalleries.Item(1).ListTemplates.Item(1),
            ContinuePreviousList=True,
            ApplyTo=win32.constants.wdListApplyToWholeList,
            DefaultListBehavior=win32.constants.wdWord10ListBehavior)
        par.Range.ListFormat.ListLevelNumber = 1

    if en_sublista or paragraph.startswith('  - '):
        par.Format.SpaceAfter = 0
        par.Range.ListFormat.ApplyListTemplateWithLevel(
            ListTemplate=par.Range.Application.ListGalleries.Item(1).ListTemplates.Item(1),
            ContinuePreviousList=True,
            ApplyTo=win32.constants.wdListApplyToWholeList,
            DefaultListBehavior=win32.constants.wdWord10ListBehavior)
        par.Range.ListFormat.ListLevelNumber = 2


# Lista para almacenar todas las tablas encontradas
all_tables_data = []

# Patrón regex para encontrar enlaces en Markdown
pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')

# Patrón regex para encontrar imágenes en Markdown rodeadas por tuberías
image1_pattern1 = re.compile(r'\|\!\[([^\]]*)\]\((.*?)\)\|')

table = None  # Añade esta línea para inicializar 'table' como None


def add_missing_delimiters(table_lines):
    new_table_lines = []
    for line in table_lines:
        stripped_line = line.rstrip()  # Elimina los espacios en blanco al final
        if stripped_line and not stripped_line.endswith("|"):
            # Si la línea no está vacía y no termina con '|', añade '|'
            stripped_line += " |"
        new_table_lines.append(stripped_line)
    return new_table_lines

while True:
    start_table = None
    end_table = None
    table_lines = []

    # Búsqueda de tablas en el documento
    for index, para in enumerate(doc.Paragraphs):
        line = para.Range.Text.strip()
        placeholders = []  # Añade esta línea para definir 'placeholders'
        
        # Verifica si la línea contiene una imagen en formato Markdown rodeada por tuberías
        if image1_pattern1.search(line):
            continue  # Si es así, ignora la línea y pasa a la siguiente

        # Ignora las líneas que comienzan con una tubería seguida de cualquier número de espacios y luego '+-' o '-+'
        if re.match(r'^\|\s*\+-', line) or re.match(r'^\|\s*-+\+', line):
            continue
        
        # Manejar inicio y fin de tabla
        if "|" in line:
            if start_table is None:
                start_table = index
            blank_line_count = 0  # Reiniciar contador de líneas vacías
            table_lines.append(line.strip())
        elif start_table is not None:
            # Considerar una línea vacía como posible fin de tabla
            if line == '':
                blank_line_count += 1
            if blank_line_count > 1:
                end_table = index
                break

    # Llama a la función para añadir delimitadores faltantes
    table_lines = add_missing_delimiters(table_lines)

    # Si no se encuentra ninguna tabla, se detiene el bucle
    if not table_lines:
        break

    # Procesamiento de la tabla en markdown
    data = []
    headers_found = False
    header_separator_line_index = None  # Índice de la línea de separación del encabezado
    for line in table_lines:
        if not headers_found and "|" in line:
            # Procesar los encabezados
            headers = [cell.strip() for cell in line.split('|')[1:-1]]
            headers = [re.sub(r':?----+:?', '', header).strip() for header in headers]  # Eliminar patrones de alineación
            if headers[-1].strip() == '':
                headers = headers[:-1]
            data.append(headers)
            headers_found = True
        elif headers_found and (re.match(r'^\|\s*-+\s*\|', line) or re.match(r'^\|\s*(:?----+:?\s*\|)+', line)):
            # Ignora la línea de separación del encabezado, no se añade a 'data'
            continue
        elif "|" in line:
            # Procesar las celdas de datos
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells[-1].strip() == '':
                cells = cells[:-1]
            cells = [cell.replace(':white_check_mark:', '✅') for cell in cells]
            data.append(cells)

    # Eliminar la fila de separación de encabezados si existe
    if header_separator_line_index is not None:
        del data[header_separator_line_index]

    counter = 0
    for para in doc.Paragraphs:
        counter += 1
        if counter == start_table:
            start_range = para.Range.Start
        if counter == end_table:
            end_range = para.Range.End
            break

    doc.Range(start_range, end_range).Delete()

    table_range = doc.Range(start_range, start_range)

    cm_to_points = 2.15 * 28.3465  # 1 cm es aproximadamente 28.3465 puntos
    
    if len(data) > 2 and len(data[0]) > 0:
        table = doc.Tables.Add(table_range, len(data), len(data[0]))
        # Configura el ancho de cada celda de la tabla
        if table is not None:
            for row in table.Rows:
                 for cell in row.Cells:
                      try:
                          cell.Width = cm_to_points
                      except Exception as e:
                              print(f"Error al ajustar el ancho de la celda: {e}")
    else:
          print("The markdown table in this bot is not recognized.")
          continue  # Continúa con el siguiente ciclo del bucle while si la tabla no se reconoce


    markdown_link_found_in_table = False  # Añade esta línea para inicializar 'markdown_link_found_in_table' en False

    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            # Comprueba si los índices están dentro del rango de la tabla
            if i < table.Rows.Count and j < table.Columns.Count:
                cell = table.Cell(i+1, j+1)
                cell_range = cell.Range             

                # Intenta ajustar el formato del párrafo
                cell_range.ParagraphFormat.ListTemplate = None
                # Ajusta el estilo del párrafo a 'Normal'
                cell_range.Style = doc.Styles('Normal')
                
                cell_range.Text = cell_data.strip()
 
                # Manipulación de hipervínculos
                matches = pattern.findall(cell_data)
                
                if matches:
                    markdown_link_found_in_table = True  # Se encontró un enlace Markdown en una celda de la tabla

                    for text, url in matches:
                        hyperlink_range = cell_range.Duplicate
                      
                        # Limpia el texto de la celda antes de añadir el hipervínculo.
                        cell_range.Text = text.strip()
                        
                        
                        hyperlink_range.Find.Execute(FindText=text)
                        doc.Hyperlinks.Add(Anchor=hyperlink_range, Address=url)
                        
                        
            else:
                print(f"Índice fuera de rango: i={i}, j={j}")

    table.Style = "Acc_Table_1"
    all_tables_data.append(data)


# Inicializa 'sheet_resized' como False
sheet_resized = False

# Recorrer las tablas y aumentar el tamaño de la hoja si hay más de 5 columnas
for table in doc.Tables:
    if table.Columns.Count > 5:
        try:
            # Aumentar el tamaño de la hoja en 2.54 cm (equivalente a 1 pulgada)
            points_in_cm = 5.45 * 28.3465  # 1 cm = 28.3465 puntos
            doc.PageSetup.PageWidth = doc.PageSetup.PageWidth + points_in_cm
            doc.PageSetup.PageHeight = doc.PageSetup.PageHeight + points_in_cm
            print(f"The size of the sheet has been increased. New width: {doc.PageSetup.PageWidth}, Nuevo alto: {doc.PageSetup.PageHeight}")
            sheet_resized = True
            break  # Salir del bucle después de la primera tabla encontrada
        except Exception as e:
            print(f"Se produjo un error al intentar ajustar el tamaño de la hoja: {e}")

if sheet_resized:
    print("The document sheet size has been resized.")
else:
    print("The size of the document sheet has not been altered.")



print("Formatting blue label b color blue /b")
# Patrón para buscar las etiquetas con diferentes colores
pattern = re.compile(r'\\(b\\)\\(a color:blue\\)(.*?)\\(/b\\)')

## Recorrer todos los párrafos del documento
for paragraph in doc.Paragraphs:
    paragraph_text = paragraph.Range.Text
    # Buscar todas las coincidencias del patrón en el párrafo
    matches = pattern.finditer(paragraph_text)
    
    for match in matches:
        matched_text = match.group(1)
        
        # Encontrar la posición del texto en el párrafo
        start = match.start(1)
        end = match.end(1)
        
        # Crear un rango para el texto
        text_range = doc.Range(paragraph.Range.Start + start, paragraph.Range.Start + end)
        
        # Aplicar el formato
        text_range.Font.Bold = True
        text_range.Font.Color = win32.constants.wdColorBlue

        # Ajustar índices para eliminar las etiquetas
        start_tag_range = doc.Range(paragraph.Range.Start + match.start(), paragraph.Range.Start + start)
        end_tag_range = doc.Range(paragraph.Range.Start + end, paragraph.Range.Start + match.end())

        # Eliminar las etiquetas
        start_tag_range.Delete()
        end_tag_range.Delete()


print("strike format word")

for para in doc.Paragraphs:
    original_text = para.Range.Text
    matches = list(re.finditer(r'~~(.*?)~~', original_text))

    # Procesar cada coincidencia en reversa para no desajustar las posiciones subsiguientes
    for match in reversed(matches):
        # Obtener las posiciones inicial y final del texto a tachar (incluyendo los símbolos ~~)
        start = match.start()
        end = match.end()

        # Obtener el texto a tachar (sin los símbolos ~~)
        text_to_strike = match.group(1)

        # Crear un rango para el texto completo a modificar (incluyendo los símbolos ~~)
        full_range = doc.Range(para.Range.Start + start, para.Range.Start + end)
        
        # Primero, aplicar el formato tachado al texto
        full_range.Font.StrikeThrough = True

        # Luego, reemplazar el texto completo (incluyendo los símbolos ~~) solo por el texto a tachar
        full_range.Text = text_to_strike

        # Actualizar el texto original para reflejar el cambio
        original_text = original_text[:start] + text_to_strike + original_text[end:]



print("Applying > signs format in the word document")
# Patrón para detectar líneas con uno o más '>'
blockquote_pattern = '^>+.*'

for para in doc.Paragraphs:
    # Verificar si el párrafo contiene el patrón de blockquote
    match = re.match(blockquote_pattern, para.Range.Text)
    if match:
        # Aplicar estilo 'Normal' antes de modificar el formato
        para.Range.Style = word_app.ActiveDocument.Styles("Normal")

        # Eliminar los caracteres '>'
        text_before_formatting = para.Range.Text
        para.Range.Text = text_before_formatting.replace('>', '', match.end() - match.start())

        # Aplicar el formato específico al párrafo
        para.Range.ParagraphFormat.Borders(constants.wdBorderLeft).LineStyle = constants.wdLineStyleSingle
        para.Range.ParagraphFormat.Borders(constants.wdBorderLeft).LineWidth = constants.wdLineWidth050pt  # Aumento del ancho
        para.Range.ParagraphFormat.Borders(constants.wdBorderLeft).Color = win32api.RGB(234, 234, 234)

        # Imprimir el texto del párrafo al que se le ha aplicado el formato
        print(f"Formato aplicado al párrafo: {text_before_formatting.strip()}")


# Función para renombrar archivos en un directorio
def rename_files_in_directory(directory):
    for filename in os.listdir(directory):
        # Dividir el nombre del archivo y la extensión
        file_root, file_extension = os.path.splitext(filename)
        # Si la extensión del archivo es .PNG o .JPG
        if file_extension.upper() in ['.PNG', '.JPG']:
            new_filename = f"{file_root}{file_extension.lower()}"
            # Renombrar el archivo
            os.rename(os.path.join(directory, filename), os.path.join(directory, new_filename))
            print(f"¡File successfully renamed to {new_filename}!")


print("etiquetas ya automaticas de colores html")

# Mapea los nombres de los colores a sus correspondientes valores RGB
colores = {
    'azul': win32api.RGB(0, 0, 255),
    'blue': win32api.RGB(0, 0, 255),
    'yellow': win32api.RGB(255, 255, 0),
    'amarillo': win32api.RGB(255, 255, 0),
    'verde': win32api.RGB(0, 128, 0),
    'green': win32api.RGB(0, 128, 0),
    'marrón': win32api.RGB(165, 42, 42),  # Definir el color marrón usando RGB
    'brown': win32api.RGB(165, 42, 42),  # Definir el color marrón usando RGB
    'rosa': win32api.RGB(255, 105, 180),
    'pink': win32api.RGB(255, 105, 180),
    'rojo': win32api.RGB(255, 0, 0),
    'red': win32api.RGB(255, 0, 0),
    'crimson': win32api.RGB(220, 20, 60),  # Color Crimson
    'teal': win32api.RGB(0, 128, 128),     # Color Teal
    'purple': win32api.RGB(128, 0, 128),   # Color Purple
    'colour': win32api.RGB(0, 0, 0),        # Color Negro para colour
    'black': win32api.RGB(0, 0, 0)
}


def aplicar_formatos(rango, texto):
    # Procesar etiquetas de color
    color_matches = re.findall(r'\(span style="color:(.*?)".*?\)(.*?)\(/span\)', texto, re.IGNORECASE)
    for match in color_matches:
        color_original = match[0]
        color = color_original.lower() 
        content = match[1]
        # Comprueba si el color está en el diccionario de colores
        if color in colores:
            try:
                rango.Font.Color = colores[color]
            except Exception as e:
                print(f"No se pudo cambiar el color: {e}")
            etiqueta_original = f'(span style="color:{color_original}")'
            texto = texto.replace(etiqueta_original, '', 1)
            texto = texto.replace('(/span)', '', 1)
        else:
            print(f"Color desconocido: {color}")

    # Procesar etiquetas de negrita
    bold_matches = re.findall(r'\(b\)(.*?)\(/b\)', texto, re.IGNORECASE)
    for match in bold_matches:
        rango.Font.Bold = constants.wdToggle
        texto = texto.replace('(b)', '').replace('(/b)', '')

    return texto

# Aplicar formatos en párrafos
for i in range(len(doc.Paragraphs), 0, -1):
    paragraph = doc.Paragraphs.Item(i)
    text = paragraph.Range.Text

    # Ignorar si no hay etiquetas HTML o si está en un campo especial
    if not re.search(r'\(span style="color:(.*?)".*?\)(.*?)\(/span\)|\(b\)(.*?)\(/b\)', text, re.IGNORECASE) or paragraph.Range.Fields.Count > 0:
        continue

    # Aplicar formatos y actualizar el texto del párrafo
    new_text = aplicar_formatos(paragraph.Range, text)
    if i <= len(doc.Paragraphs):
        doc.Paragraphs.Item(i).Range.Text = new_text

# Aplicar formatos en tablas
for table in doc.Tables:
    for row in table.Rows:
        for cell in row.Cells:
            cell_text = cell.Range.Text
            # Aplicar formatos y actualizar el texto de la celda
            new_cell_text = aplicar_formatos(cell.Range, cell_text)
            cell.Range.Text = new_cell_text


# Compilar la expresión regular para buscar imágenes en formato Markdown
image_pattern_with_pipes = re.compile(r'\|!\[([^\]]*)\]\(([^)]+)\)\|')
image_pattern_without_pipes = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')

script_dir = os.path.dirname(os.path.abspath(__file__))

# Comprobar si el directorio '.attachments' existe
attachments_dir = os.path.join(script_dir, ".attachments")
if not os.path.isdir(attachments_dir):
    print(f"Image directory does not exist: {attachments_dir}")
else:
    print(f".attachments directory exists: {attachments_dir}")

    # Renombrar los archivos en el directorio '.attachments'
    rename_files_in_directory(attachments_dir)

    # Enumerando todos los archivos en el directorio .attachments
    attachment_files = os.listdir(attachments_dir)

try:
    # Buscar y reemplazar imágenes en formato Markdown con imágenes incrustadas
    for paragraph in doc.Paragraphs:
        match_with_pipes = image_pattern_with_pipes.search(paragraph.Range.Text)
        match_without_pipes = image_pattern_without_pipes.search(paragraph.Range.Text)
        if match_with_pipes or match_without_pipes:
            if match_with_pipes:
                description = match_with_pipes.group(1)
                image_path_markdown = match_with_pipes.group(2).lstrip('/')
            elif match_without_pipes:
                description = match_without_pipes.group(1)
                image_path_markdown = match_without_pipes.group(2).lstrip('/')

            # Ignorar enlaces externos
            if image_path_markdown.startswith("http") or image_path_markdown.startswith("https"):
                continue

            # Extraer solo el nombre del archivo y la extensión del enlace de Markdown
            file_name_ext = os.path.basename(urllib.parse.unquote(image_path_markdown))

            # Comprobar si el archivo se encuentra en la lista 'attachment_files'
            if file_name_ext in attachment_files:
                image_path = os.path.join(attachments_dir, file_name_ext)

                paragraph.Range.Delete()

                image_paragraph = doc.Paragraphs.Add(paragraph.Range)
                image_paragraph.Format.Style = doc.Styles.Item("Normal")
                image_range = image_paragraph.Range
                try:
                   image_range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
                except Exception as e:
                    print(f"Se produjo un error al intentar alinear el párrafo de la imagen: {e}")

                if os.path.exists(image_path):
                    print(f"The image path exists:: {image_path}")
                    try:
                        image = image_range.InlineShapes.AddPicture(FileName=image_path, LinkToFile=False, SaveWithDocument=True)
                    except Exception as e:
                        print(f"Error al insertar la imagen: {e}")
                        continue
                else:
                    print(f"El archivo de imagen no se encuentra en la ruta especificada: {image_path}")
                    continue

                max_height = 6 * 28.3465
                if image.Height > max_height:
                    image.Height = max_height


except Exception as e:
    print(f"Capturada una excepción al procesar imágenes: {e}")
    print(f"Excepción capturada en el párrafo: {paragraph.Range.Text}")
    print(f"Detalles de la excepción: {e}")
           

# Compilar la expresión regular para buscar imágenes en formato Markdown con URL completa
image_pattern = re.compile(r'!\[([^\]]*)\]\((http[s]?://[^)]+)\)')


# Añadir una variable al inicio de tu script para llevar un seguimiento
message_shown = False

# Buscar y reemplazar imágenes en formato Markdown con imágenes incrustadas
for paragraph in doc.Paragraphs:
    match = image_pattern.search(paragraph.Range.Text)
    if match:
        # Si el mensaje aún no se ha mostrado, mostrarlo y actualizar la variable
        if not message_shown:
            print("Please wait a moment user, the script is performing very complex tasks...")
            message_shown = True

        # Obtener la descripción y la URL de la imagen
        description = match.group(1)
        image_url = match.group(2)

        paragraph.Range.Delete()

        # Crear un nuevo párrafo con estilo "Normal" para insertar la imagen
        image_paragraph = doc.Paragraphs.Add(paragraph.Range)
        image_paragraph.Format.Style = doc.Styles.Item("Normal")

        
        image_range = image_paragraph.Range
        image_range.ParagraphFormat.Alignment = win32.constants.wdAlignParagraphCenter
        image = image_range.InlineShapes.AddPicture(FileName=image_url, LinkToFile=False, SaveWithDocument=True)

        # Ajustar la altura de la imagen a un máximo de 6 cm (6 cm * 28.3465 puntos/cm = 170.078 puntos)
        max_height = 6 * 28.3465
        if image.Height > max_height:
            image.Height = max_height

        # Agregar un subtítulo a la imagen si se proporcionó una descripción
        if description:        
            # Crear un nuevo párrafo con estilo "Normal" para insertar la descripción
            desc_paragraph = doc.Paragraphs.Add(image_range)
            desc_paragraph.Range.Text = f"\\n{description}"
            desc_paragraph.Format.Alignment = win32.constants.wdAlignParagraphCenter


# Variables para detección de bloques de código
found_codeblock = False
in_codeblock = False

# Iteración sobre párrafos
for paragraph in doc.Paragraphs:
    text = paragraph.Range.Text

    # Verificar si hemos encontrado el inicio de un bloque de código
    if '```' in text:
        in_codeblock = not in_codeblock
        found_codeblock = True
        
        # Eliminar el párrafo que contiene "```"
        paragraph.Range.Delete()
        continue  # Saltar el delimitador

    # Si estamos dentro de un bloque de código, cambiar la fuente y el color del texto y del fondo
    if in_codeblock:
        paragraph.Range.Font.Name = "Consolas"
        paragraph.Range.Font.Color = win32api.RGB(255, 255, 255)  # Blanco
        paragraph.Range.Shading.BackgroundPatternColor = win32api.RGB(0, 0, 0)  # Negro

    # Si estamos dentro o fuera del bloque, eliminar los caracteres ``` antes y después.
    if '´´´' in text and not in_codeblock:  
        updated_text = text.replace('´´´', '')
        paragraph.Range.Text = updated_text


# Compilamos las expresiones regulares
markdown_pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
url_pattern = re.compile(r'-\s(http[s]?://\S+)')
markdown_link_regex = re.compile(r'-\s\[(.+?)\]\((http[s]?://\S+)\)')
markdown_pattern2 = re.compile(r'\\[([^\\]]+)\\]\(url: \\'([^\\']+)\\'\\)')

for i in range(1, doc.Paragraphs.Count + 1):
    paragraph = doc.Paragraphs.Item(i)
    match = markdown_pattern.search(paragraph.Range.Text)
    if match:
        # Comprobar si el siguiente párrafo es un título
        next_paragraph = doc.Paragraphs.Item(i + 1) if i + 1 <= doc.Paragraphs.Count else None
        if next_paragraph and next_paragraph.Style.NameLocal.startswith("Heading"):
            pass  # Puedes reemplazar 'pass' con tu propio código
        else:
            # Agrega un salto de párrafo después del enlace de Markdown
            paragraph.Range.InsertAfter('\\r')


for paragraph in doc.Paragraphs:
    match = re.search(markdown_link_regex, paragraph.Range.Text)
    if match:
        name = match.group(1)
        url = match.group(2)
        hyperlink_range = paragraph.Range
        doc.Hyperlinks.Add(hyperlink_range, url, TextToDisplay=name)
      
for paragraph in doc.Paragraphs:
    match = markdown_pattern.search(paragraph.Range.Text)
    if match:
        name = match.group(1)
        hyperlink_url = ""

        # Comprueba si el texto de anclaje del enlace de Markdown contiene una dirección web
        if re.match(r'https?://.+', match.group(2)):
            hyperlink_url = match.group(2)

        # Si el texto de anclaje del enlace de Markdown contiene una dirección web, crea un hipervínculo de Word
        if hyperlink_url:

            paragraph.Range.Delete()

            hyperlink_range = paragraph.Range
            doc.Hyperlinks.Add(hyperlink_range, hyperlink_url, TextToDisplay=name)

for paragraph in doc.Paragraphs:
    match = url_pattern.search(paragraph.Range.Text)
    if match:
        url = match.group(1)
        hyperlink_range = paragraph.Range
        # Añadir un bucle de comprobación antes de llamar al método doc.Hyperlinks.Add
        doc.Hyperlinks.Add(hyperlink_range, url)


for paragraph in doc.Paragraphs:
    match = markdown_pattern2.search(paragraph.Range.Text)
    if match:
        # Extraemos el texto del enlace y la URL
        link_text = match.group(1)
        url = match.group(2)

        # Creamos un nuevo objeto de rango que contiene el texto del enlace
        link_range = doc.Range(paragraph.Range.Start, paragraph.Range.End)
        link_range.Text = re.sub(markdown_pattern2, link_text, link_range.Text)

        # Agregamos el hipervínculo de Word
        hyperlink = doc.Hyperlinks.Add(link_range, url, TextToDisplay=link_text)

        # Cambiar el color del texto a rojo
        red_color_bgr = 255 | (0 << 8) | (0 << 16)
        hyperlink.Range.Font.Color = red_color_bgr
        hyperlink.Range.Font.Underline = True

        # Agregar una "x" al final del enlace
        hyperlink.Range.InsertAfter(' x')
      

# Definir el texto objetivo y el texto de reemplazo
target_text_1 = "###"
replacement_text_1 = ""
target_text_2 = "##"
replacement_text_2 = ""
target_text_3 = "_"
replacement_text_3 = ""
target_text_4 = "#"
replacement_text_4 = ""
target_text_5 = "*"
replacement_text_5 = ""
target_text_6 = "**"
replacement_text_6 = ""
target_text_7 = "####"
replacement_text_7 = ""
target_text_8 = "#####"
replacement_text_8 = ""


# Utilizar el objeto Find para buscar en todo el documento
find_object.ClearFormatting()

# Buscar todos los párrafos que contienen "###" y guardar su texto
found_paragraphs_1 = []
find_object.Text = target_text_1
for paragraph in doc.Paragraphs:
    if target_text_1 in paragraph.Range.Text:
        found_paragraphs_1.append(paragraph.Range.Text)

# Realizar el reemplazo de "###" por " "
if found_paragraphs_1:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_1
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_1:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_1, replacement_text_1) == paragraph.Range.Text:
            paragraph.Range.Bold = True
            paragraph.Range.Font.Size = 11
            

found_paragraphs_2 = []
find_object.Text = target_text_2
for paragraph in doc.Paragraphs:
    if target_text_2 in paragraph.Range.Text:
        found_paragraphs_2.append(paragraph.Range.Text)

# Realizar el reemplazo de "##" por " "
if found_paragraphs_2:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_2
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_2:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_2, replacement_text_2) == paragraph.Range.Text:
            paragraph.Range.Bold = True
            paragraph.Range.Font.Size = 16
           
found_paragraphs_3 = []
find_object.Text = target_text_3
for paragraph in doc.Paragraphs:
    if target_text_3 in paragraph.Range.Text:
        found_paragraphs_3.append(paragraph.Range.Text)

# Realizar el reemplazo de "_" por " "
if found_paragraphs_3:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_3
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_3:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_3, replacement_text_3) == paragraph.Range.Text:
            paragraph.Range.Italic = True
            
# Buscar todos los párrafos que contienen "#" y guardar su texto
found_paragraphs_4 = []
find_object.Text = target_text_4
for paragraph in doc.Paragraphs:
    if target_text_4 in paragraph.Range.Text:
        found_paragraphs_4.append(paragraph.Range.Text)

# Realizar el reemplazo de "#" por " "
if found_paragraphs_4:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_4
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_4:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_4, replacement_text_4) == paragraph.Range.Text:
            paragraph.Range.Bold = True
            paragraph.Range.Font.Size = 16

# Buscar todos los párrafos que contienen "**" y guardar su texto
found_paragraphs_6 = []
find_object.Text = target_text_6
for paragraph in doc.Paragraphs:
    if target_text_6 in paragraph.Range.Text:
        found_paragraphs_6.append(paragraph.Range.Text)

# Realizar el reemplazo de "**" por " "
if found_paragraphs_6:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_6
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll


for found_text in found_paragraphs_6:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_6, replacement_text_6) == paragraph.Range.Text:
            paragraph.Range.Bold = True

# Añadir el caso para "####"
found_paragraphs_7 = []
find_object.Text = target_text_7
for paragraph in doc.Paragraphs:
    if target_text_7 in paragraph.Range.Text:
        found_paragraphs_7.append(paragraph.Range.Text)

# Realizar el reemplazo de "####" por " "
if found_paragraphs_7:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_7
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_7:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_7, replacement_text_7) == paragraph.Range.Text:
            paragraph.Range.Bold = True
            paragraph.Range.Font.Size = 11


# Añadir el caso para "#####"
found_paragraphs_8 = []
find_object.Text = target_text_8
for paragraph in doc.Paragraphs:
    if target_text_8 in paragraph.Range.Text:
        found_paragraphs_8.append(paragraph.Range.Text)

# Realizar el reemplazo de "####" por " "
if found_paragraphs_8:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_8
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll

for found_text in found_paragraphs_8:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_8, replacement_text_8) == paragraph.Range.Text:
            paragraph.Range.Bold = True
            paragraph.Range.Font.Size = 11


# Buscar todos los párrafos que contienen "*" y guardar su texto
found_paragraphs_5 = []
find_object.Text = target_text_5
for paragraph in doc.Paragraphs:
    if target_text_5 in paragraph.Range.Text:
        found_paragraphs_5.append(paragraph.Range.Text)

# Realizar el reemplazo de "*" por " "
if found_paragraphs_5:
    find_object.Replacement.ClearFormatting()
    find_object.Replacement.Text = replacement_text_5
    find_object.Execute(Replace=2)  # 2 = wdReplaceAll


for found_text in found_paragraphs_5:
    for paragraph in doc.Paragraphs:
        if found_text.replace(target_text_5, replacement_text_5) == paragraph.Range.Text:
            paragraph.Range.Italic = True

print("Applying correct good formatting style if you see double dash")
for paragraph in doc.Paragraphs:
    if "---" in paragraph.Range.Text:
        # Borrar el texto original
        paragraph.Range.Text = "\\n"
        
        # Agregar una línea horizontal
        border = paragraph.Range.Borders(win32.constants.wdBorderTop)
        border.LineStyle = win32.constants.wdLineStyleSingle
        border.LineWidth = win32.constants.wdLineWidth050pt

        # Establecer el color de la línea a gris claro
        border.Color = win32api.RGB(234, 234, 234)


def apply_shading_to_text(container, pattern):
    # Iterar sobre todos los elementos en el contenedor (párrafos o celdas)
    for element in container:
        # Buscar el patrón en el texto del elemento
        matches = re.findall(pattern, element.Range.Text)
        
        # Si se encontraron coincidencias, procesar cada una
        if matches:
            for match in matches:
                # Crear un rango para la coincidencia
                start_pos = element.Range.Text.find('`' + match + '`')
                end_pos = start_pos + len(match) + 2  # Se suma 2 para incluir los caracteres de acento grave
                
                # Crear un nuevo rango que solo incluye la palabra que coincide con el patrón
                word_range = doc.Range(element.Range.Start + start_pos, element.Range.Start + end_pos)
                
                # Aplicar el sombreado gris al rango de la palabra
                word_range.Shading.BackgroundPatternColor = win32.constants.wdColorGray15
                
                # Eliminar los caracteres de acento grave
                word_range.Text = match

# Definir el patrón de búsqueda
pattern = '`(.*?)`'

print("Applying gray shading to words in paragraphs...")
apply_shading_to_text(doc.Paragraphs, pattern)

print("Applying gray shading to words in table cells...")
for table in doc.Tables:
    apply_shading_to_text(table.Range.Cells, pattern)

# Patrón regex para identificar URLs
url_pattern = re.compile(r'\b((http|https):\/\/)?[^\s()<>]+(?:\.[a-z]{2,})')

# Patrón regex para identificar imágenes en Markdown
image_pattern = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')

# Iterar sobre todos los párrafos en el documento
for paragraph in doc.Paragraphs:
    # Buscar todas las URLs en el párrafo
    for match in re.finditer(url_pattern, paragraph.Range.Text):
        # Obtener la URL y su posición en el texto del párrafo
        url = match.group()
        start_pos = match.start()
        end_pos = match.end()

         # Comprobar si la URL es parte de una imagen en Markdown
        if image_pattern.search(paragraph.Range.Text[start_pos:end_pos]):
            print(f"Ignoring the URL '{url}' because it is part of an image description in Markdown.")
            continue  # Si es así, ignora la URL y pasa a la siguiente

        # Crear un rango que solo incluye la URL
        url_range = doc.Range(paragraph.Range.Start + start_pos, paragraph.Range.Start + end_pos)

        # Añadir un hipervínculo a la URL
        doc.Hyperlinks.Add(Anchor=url_range, Address=url)  

# Itera sobre todos los campos en el documento
for field in doc.Fields:
    try:
        # Comprueba si el campo es un hipervínculo
        if field.Type == win32com.client.constants.wdFieldHyperlink:
            print("Found a hyperlink.")
            # Comprueba si el campo es parte de la tabla de contenido
            if field.Code.Text.startswith(" TOC "):
                print("The hyperlink is part of the table of contents, the formatting is not changed.")
                # Si es parte de la tabla de contenido, no cambies el formato
                continue
            # Cambia el color del texto a azul
            field.Result.Font.Color = win32com.client.constants.wdColorBlue
            # Cambia el color del subrayado a azul
            field.Result.Font.UnderlineColor = win32com.client.constants.wdColorBlue
            # Aplica un subrayado simple
            field.Result.Font.Underline = win32com.client.constants.wdUnderlineSingle
            print(f"The hyperlink format has been changed: {field.Result.Text}")
    except Exception as e:
        print(f"Se produjo un error al procesar el campo: {e}")

# Recorrer todos los párrafos del documento en orden inverso
for i in range(doc.Paragraphs.Count, 0, -1):
    paragraph = doc.Paragraphs.Item(i)
    
    # Comprobar si el párrafo contiene una imagen
    if paragraph.Range.InlineShapes.Count > 0:
        # Si el párrafo anterior es un salto de párrafo, eliminarlo
        if i > 1:  # Asegurarse de que no es el primer párrafo
            prev_paragraph = doc.Paragraphs.Item(i - 1)
            if prev_paragraph.Range.Text.strip() == "":
               prev_paragraph.Range.Delete()


# Aplicar formatos en párrafos
for i in range(len(doc.Paragraphs), 0, -1):
    paragraph = doc.Paragraphs.Item(i)
    text = paragraph.Range.Text

    # Ignorar si no hay etiquetas HTML o si está en un campo especial
    if not re.search(r'\(span style="color:(.*?)".*?\)(.*?)\(/span\)|\(b\)(.*?)\(/b\)', text, re.IGNORECASE) or paragraph.Range.Fields.Count > 0:
        continue

    # Aplicar formatos y actualizar el texto del párrafo
    new_text = aplicar_formatos(paragraph.Range, text)
    if i <= len(doc.Paragraphs):
        doc.Paragraphs.Item(i).Range.Text = new_text

# Aplicar formatos en tablas
for table in doc.Tables:
    for row in table.Rows:
        for cell in row.Cells:
            cell_text = cell.Range.Text
            # Aplicar formatos y actualizar el texto de la celda
            new_cell_text = aplicar_formatos(cell.Range, cell_text)
            cell.Range.Text = new_cell_text


print("Paragraph Cleaning...")

# Comprobar si el documento tiene al menos 6 páginas
if doc.ComputeStatistics(win32.constants.wdStatisticPages) >= 6:
    # Obtener el rango de la sexta página
    sixth_page_range = word_app.Selection.GoTo(What=win32.constants.wdGoToPage, Which=win32.constants.wdGoToAbsolute, Count=6)
else:
    print("The document is less than 6 pages.")
    sixth_page_range = None

# Recorrer todos los párrafos del documento en orden inverso
for i in range(doc.Paragraphs.Count, 0, -1):
    paragraph = doc.Paragraphs.Item(i)
    
    # Comprobar si el párrafo está en las primeras 5 páginas
    if sixth_page_range and paragraph.Range.Start < sixth_page_range.Start:
        continue  # Si está en las primeras 5 páginas, ignorarlo
    
    # Comprobar si el párrafo es un salto de párrafo
    if paragraph.Range.Text.strip() == "":
        # Si el párrafo anterior también es un salto de párrafo, eliminarlo
        if i > 1:  # Asegurarse de que no es el primer párrafo
            prev_paragraph = doc.Paragraphs.Item(i - 1)
            if prev_paragraph.Range.Text.strip() == "":
                try:
                    prev_paragraph.Range.Delete()
                except Exception as e:
                    print(f"Could not delete paragraph: {e}")

try:
    # Acceder a la tabla de contenido
    table_of_contents = doc.TablesOfContents(1)
except Exception as e:
    print(f"Se produjo un error al intentar acceder a la tabla de contenido: {e}")

try:
    # Actualizar la tabla de contenido
    table_of_contents.Update()
except Exception as e:
    print(f"Se produjo un error al intentar actualizar la tabla de contenido: {e}")

try:
    # Guardar y cerrar el documento
    doc.Save()
    doc.Close()
except Exception as e:
    print(f"Se produjo un error al intentar guardar y cerrar el documento: {e}")

try:
    word_app.Quit()
except Exception as e:
    print(f"Se produjo un error al intentar cerrar la aplicación de Word: {e}")


print("¡finished!")

    """)

# Activar el entorno virtual "shell" solo si no está activado
if 'VIRTUAL_ENV' not in os.environ:
    activate_script = os.path.join("shell", "Scripts", "activate")
    if sys.platform == "linux":
        activate_script = os.path.join("shell", "bin", "activate")

    # En Windows, necesitas ejecutar el script de activación en una shell
    if sys.platform == "win32":
        command = f"{activate_script} && {venv_python} second_script.py"
        subprocess.run(["cmd", "/k", command])
    else:  # En Unix, puedes usar 'source'
        command = f"source {activate_script} && {venv_python} second_script.py"
        subprocess.run(["bash", "-c", command])
