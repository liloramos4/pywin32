import os
import sys
import subprocess



# Comprobar si la carpeta 'shell' ya existe
if not os.path.exists('shell'):
    # Si no existe, crear un entorno virtual llamado "shell"
    subprocess.run([sys.executable, "-m", "venv", "shell"])

# Definir la ubicación del ejecutable de Python en el entorno virtual
venv_python = os.path.join("shell", "Scripts", "python")
if sys.platform == "linux":
    venv_python = os.path.join("shell", "bin", "python")

# Actualizar pip en el entorno virtual
subprocess.run([venv_python, "-m", "pip", "install", "--upgrade", "pip"])

# Instalar las dependencias especificadas en el archivo requirements.txt
subprocess.run([venv_python, "-m", "pip", "install", "-r", "requirements.txt"])

# Crear el segundo script que se ejecutará dentro del entorno virtual
with open("second_script.py", "w", encoding='utf-8') as f:
    f.write("""
# -*- coding: utf-8 -*-

import json
import base64
import win32com.client
import win32com.client as win32
import re
import win32api
import urllib.parse  # Importado para analizar la URL
import requests
import os
from docxtpl import DocxTemplate
from docx import Document
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
import docx
import time
from collections import OrderedDict
from win32com.client import constants


# Pregunta al usuario si quiere descargar todas las páginas y subpáginas de la Wiki
respuesta = input("¿Quieres descargar todas las páginas o una páginas de la Wiki? (si/no): ")

if respuesta.lower() == 'si':
    def todaslaspaginas():
        import json
        import base64
        import win32com.client
        import win32com.client as win32
        import re
        import win32api
        import urllib.parse  # Importado para analizar la URL
        import requests
        import os
        from docxtpl import DocxTemplate
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsmap
        import docx
        import time
        from collections import OrderedDict



        # Inicializa las variables globales
        stored_wiki_url = None
        stored_personal_access_token = None

        def sanitize_placeholder(placeholder):
            # Replace spaces and invalid characters with underscores
            sanitized = re.sub(r'[^\w]', '_', placeholder)
            # Remove leading digits and underscores to ensure a valid variable name
            sanitized = re.sub(r'^\d+|_', '', sanitized)
            return sanitized


        def update_toc(docx_file):
            try:
                word = win32com.client.DispatchEx("Word.Application")
                doc = word.Documents.Open(docx_file)
                doc.TablesOfContents(1).Update()
                doc.Close(SaveChanges=True)
                word.Quit()
            except Exception as e:
                print(f"An error occurred while updating the table of contents: {e}")


        def get_page_content(url):
            content_url = url + "?api-version=7.0&includeContent=true"
            response = requests.get(content_url, headers=headers)

            if response.status_code == 200:
                return json.loads(response.text)['content']
            else:
                print(f"Error al obtener el contenido de la página: {response.status_code}")
                return "Contenido no disponible"

        def extract_pages_recursive(page, level=1):
            if not page:
                return []

            original_content = get_page_content(page['url'])
            content = get_page_content(page['url'])
            # Replace <span> tags with plain text
            content = re.sub(r'(<span style="color:([^>]*)">([^<]*?))\\n', r'\\1(/span)\\n', content)
            # Reemplazar <center>(.*?)</center> con '(.*?)'
            content = re.sub(r'<center>(.*?)</center>', r'\\1', content)
            # reemplazar ciertos bloques códigos wiki
            content = re.sub(r'<code>```[^\\n]*\\n', '```\\n', content)
            # Reemplazar <br>(.*?)</br> con '(.*?)'
            content = re.sub(r'<br>(.*?)</br>', r'\\1', content)
            # reemplaza etiqueta <br> con un salto de linea
            content = re.sub(r'<br>', '\\n', content)
            # Reemplazar To_do @<lo que sea> con TO_DO @lo que sea
            content = re.sub(r'TO_DO: @<([A-Fa-f0-9-]+)>', r'TO_DO: \\1', content)
            # Reemplazar <Lista> @<lo que sea> con Lista @lo que sea
            content = re.sub(r'<Lista> @<([^>]+)>', r'Lista @\\1', content)
            # quitar mayusculas formatos de imagen
            content = re.sub(r'\\.(PNG|JPG|JPEG|GIF)', lambda x: x.group().lower(), content)
            content = re.sub(r'<([^>]*)>', r'(\\1)', content)
            # Reemplazar '(/code)' con ''
            content = re.sub(r'\(/code\\)', '', content)
            

            info = {
                'name': page['path'],
                'short_name': page['path'].split('/')[-1],
                'url': page['url'],
                'original_content': original_content,
                'content': content ,
                'level': level,
                'subpages': []
            }


            page_info = [info]

            if 'subPages' in page:
                for sub_page in page['subPages']:
                    page_info.extend(extract_pages_recursive(sub_page, level+1))

            return page_info

        def extract_url_values(url):
            regex = r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/.*"
            match = re.search(regex, url)
            if match:
                return match.groupdict()
            else:
                return None



        def extract_placeholders(template_path):
            doc = Document(template_path)
            placeholders = []

            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        
                        # Find placeholders in the full text
                        start_index = full_text.find('{{')
                        end_index = full_text.find('}}')
                        while start_index != -1 and end_index != -1:
                            placeholder = full_text[start_index+2:end_index].strip()
                            placeholders.append(placeholder)
                            
                            # Find the next placeholder
                            start_index = full_text.find('{{', end_index)
                            end_index = full_text.find('}}', end_index+2)

                elif element.tag.endswith('tc'):  # Check for table cell (td)
                    for p in element.iterchildren('{%s}p' % nsmap['w']):
                        paragraph = Paragraph(p, doc)
                        if hasattr(paragraph, 'runs'):
                            # Concatenate the text of adjacent runs
                            full_text = ''.join([run.text for run in paragraph.runs])
                            
                            # Find placeholders in the full text
                            start_index = full_text.find('{{')
                            end_index = full_text.find('}}')
                            while start_index != -1 and end_index != -1:
                                placeholder = full_text[start_index+2:end_index].strip()
                                placeholders.append(placeholder)
                                
                                # Find the next placeholder
                                start_index = full_text.find('{{', end_index)
                                end_index = full_text.find('}}', end_index+2)

            # Remove duplicates from the list of placeholders while maintaining the order of the elements
            placeholders = list(OrderedDict.fromkeys(placeholders))

            return placeholders


        def create_context(page_info, placeholders):
            context = {}
            title_index = 1
            for placeholder in placeholders:
                if title_index < len(page_info):
                    page = page_info[title_index]
                    # Agregar valores al contexto tanto para marcadores de posición de título como para marcadores de posición de contenido
                    if placeholder.endswith('_content'):
                        if page['content'] != 'No hay contenido':
                            context[placeholder] = page['content']
                        title_index += 1
                    else:
                        context[placeholder] = page['short_name']
            return context


        # Pedir la información al usuario una sola vez y almacenarla en variables globales
        # Verificar si las credenciales ya se han almacenado
        if stored_wiki_url is None or stored_personal_access_token is None:
            wiki_url = input("Introduce la URL principal del portal Wiki de Azure: ")
            personal_access_token = input("Introduce tu token de acceso personal: ")
            # Aquí iría la lógica para verificar si las credenciales son válidas.
            # Si son válidas, las almacenamos en las variables globales.
            stored_wiki_url = wiki_url
            stored_personal_access_token = personal_access_token
        else:
            # Usar las credenciales almacenadas
            wiki_url = stored_wiki_url
            personal_access_token = stored_personal_access_token


        credentials = f":{personal_access_token}"
        encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')

        headers = {
            'Authorization': f'Basic {encoded_credentials}',
            "Content-Type": "application/json",
            "Accept": "application/json"
        }

        # Decodifica la URL para manejar caracteres especiales
        decoded_wiki_url = urllib.parse.unquote(wiki_url)
        # Usa regex para extraer la organización, el proyecto y el wiki de la URL decodificada
        url_match = re.search(r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/.*", decoded_wiki_url)
        if url_match:
            organization = url_match.group('organization')
            project = url_match.group('project')
            wiki = url_match.group('wiki')
        else:
            print("La URL proporcionada no es válida.")
            exit()

        url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki}/pages?api-version=7.0&recursionLevel=full"

        response = requests.get(url, headers=headers)

        if response.status_code == 200:
            print("Respuesta exitosa http 200.")
            root_page = json.loads(response.text)
        else:
            print(f"Error al obtener la página raíz de la wiki: {response.status_code}")
            root_page = {}

        page_info = extract_pages_recursive(root_page)


        # Save the .md file with original content
        md_filename_original = 'htmlymd.md'
        with open(md_filename_original, 'w', encoding='utf-8') as f:
            for page in page_info:
                # Get the title and original content of the page
                title = page['name'].split('/')[-1]
                original_content = page['original_content'].strip()

                # Only write the title and original content to the Markdown file if the original content is not empty
                if original_content and not (title.startswith("#") and "No hay contenido" in title) and page['name'] != "/":
                    f.write(f'# {title}\\n')
                    f.write(original_content)
                    f.write('\\n\\n')


        # Save the .md file
        md_filename = 'todosmd.md'
        with open('todosmd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Get the title and content of the page
                title = page['name'].split('/')[-1]
                content = page['content'].strip()

                # Only write the title and content to the Markdown file if the content is not empty
                if content and not (title.startswith("#") and "No hay contenido" in title) and page['name'] != "/":
                    f.write(f'# {title}\\n')
                    f.write(content)
                    f.write('\\n\\n')


        # Read the Markdown content
        with open(md_filename, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        # Extract the level 1 headings with regex
        title_pattern = r'^# (.*)'
        titles = re.findall(title_pattern, markdown_content, flags=re.M)


        # Generate the placeholders
        placeholders = [{'text': f'Titulo{i+1}', 'level': page['level']} for i, page in enumerate(page_info)]

        def add_titles_to_template(template_path, titles_info):
            doc = Document(template_path)

            # Find the element after which to insert the placeholders
            insert_element = None
            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if 'REFERENCES' in full_text:
                            insert_element = paragraph
                            break

            if insert_element is not None:
                # Insert the titles before the element with the desired heading style
                for title_info in titles_info:
                    # Skip titles with name '/' or empty short_name
                    if 'name' in title_info and (title_info['name'] == '/' or not title_info['short_name']):
                        continue
                    sanitized_title = sanitize_placeholder(title_info["title"])
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}}}}}')
                    # Calculate the maximum level in the JSON
                    max_level = max([title_info['level'] for title_info in titles_info])
                    # Generate the level_map dictionary dynamically
                    level_map = {i: f'Heading {i}' for i in range(1, max_level+1)}
                    heading_style = level_map.get(title_info["level"] - 1, 'Normal')
                    p.style = heading_style
                    # Insert a placeholder for the content below the title
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}_content}}}}')

                # Insert a page break after the last title
                p = insert_element.insert_paragraph_before()
                run = p.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)

            # Save the modified template
            doc.save('new_template.docx')


        # Generate the title information
        titles_info = [{'title': page['name'].split('/')[-1], 'level': page['level']} for page in page_info if page['content'].strip() and page['name'].split('/')[-1].strip()]


        # Call the function to create a new template and add the titles to it
        add_titles_to_template('plantilla.docx', titles_info)

        # Load the modified template
        doc = DocxTemplate("new_template.docx")

        # Extract the placeholders from the template
        placeholders = extract_placeholders("new_template.docx")


        # Create a new context
        context = create_context(page_info, placeholders)


        # Render template with dynamic context
        doc.render(context)

        # Save the generated document
        doc.save("documento_generado.docx")


        print("Wait one moment to finish work")
        # Wait for a few seconds to make sure the file is saved
        time.sleep(4)
        print(" already ")

        # Check if the file was saved successfully
        if os.path.exists("documento_generado.docx"):
            print("El archivo documento_generado.docx se ha guardado correctamente.")
        else:
            print("No se pudo guardar el archivo documento_generado.docx.")

        # Get the directory of the script
        dir_path = os.path.dirname(os.path.realpath(__file__))

        # Build the file path
        docx_file = os.path.join(dir_path, "documento_generado.docx")

        # Call the function to update the table of contents in the generated document
        update_toc(docx_file)  

        pass
        
    todaslaspaginas()  # Llama a la función todaslaspaginas

else:

    def paginaconcreta2():
        import requests
        import json
        import base64
        import re
        from docxtpl import DocxTemplate
        import os
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsmap
        from collections import OrderedDict
        import docx
        import win32com.client
        import time
        import urllib.parse

       
        # Inicializa las variables globales
        stored_wiki_url = None
        stored_personal_access_token = None

        def sanitize_placeholder(placeholder):
            # Replace spaces and invalid characters with underscores
            sanitized = re.sub(r'[^\w]', '_', placeholder)
            # Remove leading digits and underscores to ensure a valid variable name
            sanitized = re.sub(r'^\d+|_', '', sanitized)
            return sanitized

        def update_toc(docx_file):
            word = win32com.client.DispatchEx("Word.Application")
            doc = word.Documents.Open(docx_file)
            doc.TablesOfContents(1).Update()
            doc.Close(SaveChanges=True)
            word.Quit()


        def get_page_content(url, headers):
            content_url = url + "?api-version=7.0&includeContent=true"
            response = requests.get(content_url, headers=headers)

            if response.status_code == 200:
                # Convierte la respuesta en un objeto JSON
                response_json = json.loads(response.text)
                # Devuelve el valor del campo 'content'
                return response_json['content']
            else:
                print(f"Error al obtener el contenido de la página: {response.status_code}")
                return "Contenido no disponible"

        def extract_pages_recursive(page, headers=None, level=1):
            if not page:
                return []

            original_content = get_page_content(page['url'], headers)
            content = get_page_content(page['url'], headers)
            # Replace <span> tags with plain text
            content = re.sub(r'(<span style="color:([^>]*)">([^<]*?))\\n', r'\\1(/span)\\n', content)
            # Reemplazar <center>(.*?)</center> con '(.*?)'
            content = re.sub(r'<center>(.*?)</center>', r'\\1', content)
            # reemplazar ciertos bloques códigos wiki
            content = re.sub(r'<code>```[^\\n]*\\n', '```\\n', content)
            # Reemplazar <br>(.*?)</br> con '(.*?)'
            content = re.sub(r'<br>(.*?)</br>', r'\\1', content)
            # reemplaza etiqueta <br> con un salto de linea
            content = re.sub(r'<br>', '\\n', content)
            # Reemplazar To_do @<lo que sea> con TO_DO @lo que sea
            content = re.sub(r'TO_DO: @<([A-Fa-f0-9-]+)>', r'TO_DO: \\1', content)
            # Reemplazar <Lista> @<lo que sea> con Lista @lo que sea
            content = re.sub(r'<Lista> @<([^>]+)>', r'Lista @\\1', content)
            # quitar mayusculas formatos de imagen
            content = re.sub(r'\\.(PNG|JPG|JPEG|GIF)', lambda x: x.group().lower(), content)
            content = re.sub(r'<([^>]*)>', r'(\\1)', content)
            # Reemplazar '(/code)' con ''
            content = re.sub(r'\(/code\\)', '', content)
            
                
                   
            info = {
                'name': page['path'],
                'short_name': page['path'].split('/')[-1],
                'url': page['url'],
                'original_content': original_content,
                'content': content,
                'level': level,
                'subpages': []
            }

            page_info = [info]

            if 'subPages' in page:
                for sub_page in page['subPages']:
                    page_info.extend(extract_pages_recursive(sub_page, headers, level+1))

            return page_info

        def extract_url_values(url):
            regex = r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>[^/]+)/(?P<page_id>\d+)/.*"
            match = re.search(regex, url)
            if match:
                return match.groupdict()
            else:
                return None

      
        def obtain_page(organization, project, wiki, page_id, headers):
            api_url = f"https://dev.azure.com/{organization}/{project}/_apis/wiki/wikis/{wiki}/pages/{page_id}?api-version=7.0&recursionLevel=full&includeContent=true"
            response = requests.get(api_url, headers=headers)
            if response.status_code == 200:
                print("Respuesta exitosa http 200.")
                page = json.loads(response.text)
                return page
            else:
                print(f"Error al obtener la página de la Wiki: {response.status_code}")

            return None
    

        def download_specific_page(headers, wiki_url2):
            url_values = extract_url_values(wiki_url2)

            if url_values is None:
                print("La URL proporcionada no es válida.")
                return

            page_id = url_values['page_id']
            organization = url_values['organization']
            project = url_values['project']
            wiki = url_values['wiki']

            page = obtain_page(organization, project, wiki, page_id, headers)

            # Añade una condición de parada para la recursión
            if page is not None:
                return organization, project, wiki, page
            else:
                print("No se pudo descargar la página.")
                return None


                
        # Pedir la información al usuario una sola vez y almacenarla en variables globales
        # Verificar si las credenciales ya se han almacenado
        if stored_wiki_url is None or stored_personal_access_token is None or stored_wiki_url2 is None:
            wiki_url2 = input("Introduce la URL de la página que quieres buscar: ")
            personal_access_token = input("Introduce tu token de acceso personal: ")
            print("Por favor usuario necesito que especifiques tu URL principal azure wiki para poder descargar los ficheros tanto fotos u otros ficheros")
            wiki_url = input("Introduce la URL principal del portal Wiki de Azure: ")
            # Aquí iría la lógica para verificar si las credenciales son válidas.
            # Si son válidas, las almacenamos en las variables globales
            stored_wiki_url = wiki_url
            stored_personal_access_token = personal_access_token
        else:
            # Usar las credenciales almacenadas
            wiki_url = stored_wiki_url
            personal_access_token = stored_personal_access_token

        
        credentials = f":{personal_access_token}"
        encoded_credentials = base64.b64encode(credentials.encode('utf-8')).decode('utf-8')
        headers = {
                'Authorization': f'Basic {encoded_credentials}',
                "Content-Type": "application/json",
                "Accept": "application/json"
        }

        result = download_specific_page(headers, wiki_url2)
        if result is not None:
              organization, project, wiki, page_data = result
              page_info = extract_pages_recursive(page_data, headers)  # Asegúrate de que esto devuelva una lista 
        else:
            print("No se pudo descargar la página.")
            

        def extract_url_values(wiki_url2):
            # Decodifica la URL para manejar caracteres especiales
            decoded_wiki_url2 = urllib.parse.unquote(wiki_url2)
            # Usa regex para extraer la organización, el proyecto y el wiki de la URL decodificada
            api_url_match = re.search(r"https://dev.azure.com/(?P<organization>[^/]+)/(?P<project>[^/]+)/_wiki/wikis/(?P<wiki>.+).wiki/(?P<page_id>\d+)/.*", decoded_wiki_url2)
            if api_url_match:
                 organization = api_url_match.group('organization')
                 project = api_url_match.group('project')
                 wiki = api_url_match.group('wiki')
                 page_id = api_url_match.group('page_id')
                 return {'organization': organization, 'project': project, 'wiki': wiki, 'page_id': page_id}
            else:
                 print("La URL proporcionada no es válida.")
                 exit() 

        
        # Escribir el contenido original de las páginas en un archivo Markdown
        with open('htmlymd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Obtener el título y contenido original de la página
                title = page['name'].split('/')[-1]
                original_content = page['original_content'].strip()

                # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                if original_content:
                    f.write(f'# {title}\\n')
                    f.write(original_content)
                    f.write('\\n\\n')

                # Recorrer las subpáginas y escribir su contenido original
                for subpage in page['subpages']:
                    # Obtener el título y contenido original de la subpágina
                    subpage_title = subpage['name'].split('/')[-1]
                    subpage_original_content = subpage['original_content'].strip()

                    # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                    if subpage_original_content:
                        f.write(f'# {subpage_title}\\n')
                        f.write(subpage_original_content)
                        f.write('\\n\\n')

                
        # Escribir el contenido de las páginas en un archivo Markdown
        with open('todosmd.md', 'w', encoding='utf-8') as f:
            for page in page_info:
                # Obtener el título y contenido de la página
                title = page['name'].split('/')[-1]
                content = page['content'].strip()
                

                # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                if content:
                    f.write(f'# {title}\\n')
                    f.write(content)
                    f.write('\\n\\n')

                # Recorrer las subpáginas y escribir su contenido
                for subpage in page['subpages']:
                    # Obtener el título y contenido de la subpágina
                    subpage_title = subpage['name'].split('/')[-1]
                    subpage_content = subpage['content'].strip()
                   
                    # Solo escribir el título y contenido en el archivo Markdown si el contenido no está vacío
                    if subpage_content:
                        f.write(f'# {subpage_title}\\n')
                        f.write(subpage_content)
                        f.write('\\n\\n')
                        
                else:
                    if not page['content'].strip():
                            print(f"No se encontró la página con ID {page_id}.")


        def extract_placeholders(template_path):
            doc = Document(template_path)
            placeholders = []

            previous_full_text = None  # Agregar variable para almacenar el texto completo anterior

            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if full_text != previous_full_text:  # Solo imprimir el texto completo si es diferente del anterior
                            previous_full_text = full_text  # Actualizar el valor de previous_full_text
                        
                        # Find placeholders in the full text
                        start_index = full_text.find('{{')
                        end_index = full_text.find('}}')
                        while start_index != -1 and end_index != -1:
                            placeholder = full_text[start_index+2:end_index].strip()
                            placeholders.append(placeholder)
                            
                            # Find the next placeholder
                            start_index = full_text.find('{{', end_index)
                            end_index = full_text.find('}}', end_index+2)

                elif element.tag.endswith('tc'):  # Check for table cell (td)
                    for p in element.iterchildren('{%s}p' % nsmap['w']):
                        paragraph = Paragraph(p, doc)
                        if hasattr(paragraph, 'runs'):
                            # Concatenate the text of adjacent runs
                            full_text = ''.join([run.text for run in paragraph.runs])
                            if full_text != previous_full_text:  # Solo imprimir el texto completo si es diferente del anterior
                                previous_full_text = full_text  # Actualizar el valor de previous_full_text
                            
                            # Find placeholders in the full text
                            start_index = full_text.find('{{')
                            end_index = full_text.find('}}')
                            while start_index != -1 and end_index != -1:
                                placeholder = full_text[start_index+2:end_index].strip()
                                placeholders.append(placeholder)
                                
                                # Find the next placeholder
                                start_index = full_text.find('{{', end_index)
                                end_index = full_text.find('}}', end_index+2)

            # Remove duplicates from the list of placeholders while maintaining the order of the elements
            placeholders = list(OrderedDict.fromkeys(placeholders))

            return placeholders

        def create_context(page_info, placeholders):
            context = {}
            title_index = 0
            for placeholder in placeholders:
                if title_index < len(page_info):
                    page = page_info[title_index]
                    # Agregar valores al contexto tanto para marcadores de posición de título como para marcadores de posición de contenido
                    if placeholder.endswith('_content'):
                        # Eliminar espacios adicionales del contenido
                        content = re.sub(r'\s+', ' ', page['content'])
                        # Eliminar las dos almohadillas del contenido
                        content = content.replace('##', '')
                        context[placeholder] = page['content']
                        title_index += 1
                    else:
                        context[placeholder] = page['name'].split('/')[-1]
            # Imprimir el contexto para depuración
            print("Contexto inicial original:", context)
            return context


        def add_titles_to_template(template_path, titles_info):
            doc = Document(template_path)

            # Find the element after which to insert the placeholders
            insert_element = None
            for element in doc.element.body.iter():
                if element.tag.endswith('p'):
                    paragraph = Paragraph(element, doc)
                    if hasattr(paragraph, 'runs'):
                        # Concatenate the text of adjacent runs
                        full_text = ''.join([run.text for run in paragraph.runs])
                        if 'REFERENCES' in full_text:
                            insert_element = paragraph
                            break

            if insert_element is not None:
                # Insert the titles before the element with the desired heading style
                for title_info in titles_info:
                    sanitized_title = sanitize_placeholder(title_info["title"])
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}}}}}')
                    # Map levels to Word Heading styles
                    level_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4', 5: 'Heading 5'}
                    heading_style = level_map.get(title_info["level"], 'Normal')
                    p.style = heading_style

                    # Insert a placeholder for the content below the title
                    p = insert_element.insert_paragraph_before(f'{{{{{sanitized_title}_content}}}}')
                    # Set the alignment of the paragraph to left
                    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Apply bold formatting to the text if it starts with two hash symbols and remove the hash symbols from the content
                    if p.text.startswith('##'):
                        for run in p.runs:
                            run.bold = True
                            run.text = run.text.replace('##', '')

                # Insert a page break after the last title
                p = insert_element.insert_paragraph_before()
                run = p.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)

            # Save the modified template
            doc.save('new_template.docx')


        # Generate the title information
        titles_info = [{'title': page['name'].split('/')[-1], 'level': page['level']} for page in page_info if page['content'].strip() and page['name'].split('/')[-1].strip()]


        # Call the function to create a new template and add the titles to it
        add_titles_to_template('plantilla.docx', titles_info)

        # Load the modified template
        doc = DocxTemplate("new_template.docx")

        # Extract the placeholders from the template
        placeholders = extract_placeholders("new_template.docx")

        # Create a new context
        context = create_context(page_info, placeholders)

        # Imprimir el contexto para depuración resultado final
        print("Contexto creado finalizado:", context)

        # Render template with dynamic context
        doc.render(context)

        # Save the generated document
        doc.save("documento_generado.docx")


        print("Wait one moment to finish work")
        # Wait for a few seconds to make sure the file is saved
        time.sleep(4)
        print("already")


        # Check if the file was saved successfully   
        if os.path.exists("documento_generado.docx"):
            print("El archivo documento_generado.docx se ha guardado correctamente.")
        else:
            print("No se pudo guardar el archivo documento_generado.docx.")

        # Get the directory of the script
        dir_path = os.path.dirname(os.path.realpath(__file__))

        # Build the file path
        docx_file = os.path.join(dir_path, "documento_generado.docx")

        # Call the function to update the table of contents in the generated document
        update_toc(docx_file)

        document = Document('documento_generado.docx')



    paginaconcreta2()  # Llama a la función paginaconcreta2



print("¡Finalizado!")
    """)

# Activar el entorno virtual "shell" solo si no está activado
if 'VIRTUAL_ENV' not in os.environ:
    activate_script = os.path.join("shell", "Scripts", "activate")
    if sys.platform == "linux":
        activate_script = os.path.join("shell", "bin", "activate")

    # En Windows, necesitas ejecutar el script de activación en una shell
    if sys.platform == "win32":
        command = f"{activate_script} && {venv_python} second_script.py"
        subprocess.run(["cmd", "/k", command])
    else:  # En Unix, puedes usar 'source'
        command = f"source {activate_script} && {venv_python} second_script.py"
        subprocess.run(["bash", "-c", command])
